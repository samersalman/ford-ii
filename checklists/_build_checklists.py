"""
FORD-II v12 — Checklist Build Script

Reads the three blank reporting-guideline templates (CHEERS-2022, STROBE,
TRIPOD) and the three map modules (_cheers_map, _strobe_map, _tripod_map),
fills the destination cells, and emits 6 outputs:

    CHEERS-2022-checklist-completed.docx
    CHEERS-2022-checklist-completed.md
    STROBE-checklist-v4-combined-PlosMedicine-completed.docx
    STROBE-checklist-v4-combined-PlosMedicine-completed.md
    Tripod-Checklist-Prediction-Model-Development-and-Validation-Word-completed.docx
    Tripod-Checklist-Prediction-Model-Development-and-Validation-Word-completed.md

Run from this directory:

    python _build_checklists.py

The script is idempotent: re-running overwrites the previous outputs with
byte-identical content (modulo docx zip ordering / mtime, which python-docx
preserves on a clean rewrite).
"""

from __future__ import annotations

import datetime as _dt
import json
import os
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

HERE = Path(__file__).resolve().parent
PARENT = HERE.parent  # …/v12/checklists

CHEERS_TEMPLATE = PARENT / "CHEERS-2022-checklist.docx"
STROBE_TEMPLATE = PARENT / "STROBE-checklist-v4-combined-PlosMedicine.docx"
TRIPOD_TEMPLATE = (
    PARENT / "Tripod-Checklist-Prediction-Model-Development-and-Validation-Word.docx"
)

CHEERS_OUT_DOCX = HERE / "CHEERS-2022-checklist-completed.docx"
CHEERS_OUT_MD = HERE / "CHEERS-2022-checklist-completed.md"
STROBE_OUT_DOCX = HERE / "STROBE-checklist-v4-combined-PlosMedicine-completed.docx"
STROBE_OUT_MD = HERE / "STROBE-checklist-v4-combined-PlosMedicine-completed.md"
TRIPOD_OUT_DOCX = (
    HERE
    / "Tripod-Checklist-Prediction-Model-Development-and-Validation-Word-completed.docx"
)
TRIPOD_OUT_MD = (
    HERE
    / "Tripod-Checklist-Prediction-Model-Development-and-Validation-Word-completed.md"
)

MANUSCRIPT_TITLE = (
    "Fracture Orthopedic Risk of Non-Home Discharge II (FORD-II): "
    "National Registry Derivation, Internal Validation, and "
    "Cost-Effectiveness of a Fracture Trauma Discharge-Disposition Score"
)
GEN_DATE = _dt.date.today().isoformat()


# ---------------------------------------------------------------------------
# Map loader
# ---------------------------------------------------------------------------


def load_maps() -> dict[str, dict[str, Any]]:
    """Import the three map modules and return their dicts."""
    sys.path.insert(0, str(HERE))
    from _cheers_map import CHEERS_MAP  # type: ignore  # noqa: E402
    from _strobe_map import STROBE_MAP  # type: ignore  # noqa: E402
    from _tripod_map import TRIPOD_MAP  # type: ignore  # noqa: E402

    return {"cheers": CHEERS_MAP, "strobe": STROBE_MAP, "tripod": TRIPOD_MAP}


# ---------------------------------------------------------------------------
# Cell-writing primitives
# ---------------------------------------------------------------------------


def _clear_cell(cell) -> None:
    """Remove all paragraph elements from a cell, leaving exactly one empty one.

    python-docx requires a cell to retain at least one <w:p> element, so we
    delete all <w:p> children and append a fresh empty paragraph.
    """
    tc = cell._tc
    # Remove every paragraph child (w:p)
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    # Also remove any straggling tables-in-cell (not expected, but safe)
    for tbl in tc.findall(qn("w:tbl")):
        tc.remove(tbl)
    # Append a single empty paragraph so the cell is valid XML
    cell.add_paragraph()
    # The newly added paragraph is the only one; we'll write text via add_paragraph
    # later, so drop the empty placeholder by replacing on first write.


def write_cell_paragraphs(cell, lines: list[str]) -> None:
    """Replace the cell's contents with the given paragraphs (1 per line).

    Preserves default paragraph style. We do NOT set runs to override fonts.
    """
    _clear_cell(cell)
    # _clear_cell left exactly one empty paragraph; use it for the first line.
    first_para = cell.paragraphs[0]
    if not lines:
        return
    first_para.text = lines[0]
    for line in lines[1:]:
        cell.add_paragraph(line)


def format_page_line(entry: dict[str, Any]) -> str:
    """Build 'p. {page} — {section_ref}' or 'N/A — {reason}'."""
    if entry.get("na_reason"):
        return f"N/A — {entry['na_reason']}"
    page = entry.get("page")
    section = entry.get("section_ref") or ""
    if page is None:
        return f"N/A — {section}".rstrip(" —")
    return f"p. {page} — {section}".rstrip(" —")


def format_quote_line(entry: dict[str, Any]) -> str | None:
    """Build the quote paragraph, or None if N/A (no quote)."""
    if entry.get("na_reason"):
        return None
    quote = entry.get("quote")
    if not quote:
        return None
    return f"“{quote}”"


# ---------------------------------------------------------------------------
# Item-id extraction
# ---------------------------------------------------------------------------

ITEM_RE = re.compile(r"\d+[a-z]?")  # e.g. 1, 1a, 12, 12d, 13a


def extract_item_id(raw: str) -> str | None:
    """Pull a canonical item id (e.g. '1a', '12', '16c') from a template cell.

    Strips whitespace, slashes, asterisks, etc. Returns None if no digit
    sequence is found (header / section-band rows).
    """
    if raw is None:
        return None
    txt = raw.strip().lower()
    if not txt:
        return None
    # Discard rows whose item cell is a section label (e.g. "Introduction")
    if not any(ch.isdigit() for ch in txt):
        return None
    # Strip leading "/" or whitespace fragments
    txt = txt.replace("/", " ")
    m = ITEM_RE.search(txt)
    return m.group(0) if m else None


def extract_item_from_row_cohort(row, recommendation_text: str) -> str | None:
    """For STROBE, pull item id from col 1 and the (a)/(b)/(c) suffix from col 2.

    STROBE template repeats item number '12' across 12a-12e rows; the actual
    sub-letter is encoded as '(a)' '(b)' '(c)' inside the recommendation cell.
    """
    base = extract_item_id(row.cells[1].text)
    if base is None:
        return None
    # Look for "(a)", "(b)", etc. in the recommendation cell
    m = re.match(r"\s*\(([a-e])\)", recommendation_text)
    if m:
        return f"{base}{m.group(1)}"
    return base


# ---------------------------------------------------------------------------
# CHEERS filler
# ---------------------------------------------------------------------------


def fill_cheers_docx(map_dict: dict[str, dict], in_path: Path, out_path: Path) -> dict:
    """Fill the CHEERS Reported-in-section column (col 3).

    Returns coverage stats.
    """
    doc = Document(str(in_path))
    table = doc.tables[0]
    matched, unmapped, na_count, filled_cells = 0, [], 0, 0
    expected_keys = set(map_dict.keys())
    seen_keys: set[str] = set()

    for ri, row in enumerate(table.rows):
        if ri == 0:
            continue  # header
        item_id = extract_item_id(row.cells[1].text)
        if item_id is None:
            continue  # section band (TITLE, ABSTRACT, etc.)
        entry = map_dict.get(item_id)
        if entry is None:
            unmapped.append((ri, row.cells[1].text.strip()))
            print(
                f"[CHEERS] WARNING: row {ri} item '{item_id}' not in map",
                file=sys.stderr,
            )
            continue
        seen_keys.add(item_id)
        matched += 1
        lines = []
        page_line = format_page_line(entry)
        lines.append(page_line)
        q = format_quote_line(entry)
        if q is not None:
            lines.append(q)
        if entry.get("na_reason"):
            na_count += 1
        dest = row.cells[3]
        write_cell_paragraphs(dest, lines)
        filled_cells += 1

    doc.save(str(out_path))
    orphans = sorted(expected_keys - seen_keys)
    return {
        "matched": matched,
        "expected": len(expected_keys),
        "unmapped_rows": unmapped,
        "orphan_keys": orphans,
        "na_count": na_count,
        "filled_cells": filled_cells,
        "n_tables": len(doc.tables),
        "n_rows_by_table": [len(t.rows) for t in doc.tables],
    }


# ---------------------------------------------------------------------------
# STROBE filler
# ---------------------------------------------------------------------------


def fill_strobe_docx(map_dict: dict[str, dict], in_path: Path, out_path: Path) -> dict:
    """Fill STROBE Page (col 3) and Relevant text (col 4) across 3 tables."""
    doc = Document(str(in_path))
    if len(doc.tables) != 3:
        raise SystemExit(
            f"[STROBE] FATAL: expected 3 tables, got {len(doc.tables)} — "
            f"template structure has changed."
        )

    matched, unmapped, na_count, filled_cells = 0, [], 0, 0
    table_stats: list[dict[str, int]] = []
    expected_keys = set(map_dict.keys())
    seen_keys: set[str] = set()

    # Item 15 in STROBE has 3 design-variant rows (cohort, case-control,
    # cross-sectional). Only the cohort row gets filled. Track which rows
    # mention 'Cohort study—' in their recommendation cell.
    for ti, table in enumerate(doc.tables):
        t_match = t_skip = 0
        for ri, row in enumerate(table.rows):
            # Section bands: col 0 == col 1 == col 2 == col 3 (template treats them
            # as merged header rows). Detect by checking that col 1 has no digit.
            item_no_text = row.cells[1].text.strip()
            if not any(ch.isdigit() for ch in item_no_text):
                continue
            rec_text = row.cells[2].text
            # Item 15: skip case-control / cross-sectional variant rows
            if item_no_text.startswith("15"):
                if "Case-control study" in rec_text or "Cross-sectional study" in rec_text:
                    t_skip += 1
                    continue
            item_id = extract_item_from_row_cohort(row, rec_text)
            if item_id is None:
                continue
            entry = map_dict.get(item_id)
            if entry is None:
                unmapped.append((ti, ri, item_id))
                print(
                    f"[STROBE] WARNING: table {ti} row {ri} item '{item_id}' "
                    f"not in map",
                    file=sys.stderr,
                )
                continue
            seen_keys.add(item_id)
            matched += 1
            t_match += 1

            # Col 3 = Page No.
            page_cell = row.cells[3]
            if entry.get("na_reason"):
                write_cell_paragraphs(page_cell, ["N/A"])
                na_count += 1
            else:
                page = entry.get("page")
                write_cell_paragraphs(
                    page_cell, [f"p. {page}" if page is not None else "N/A"]
                )
            filled_cells += 1

            # Col 4 = Relevant text from manuscript
            text_cell = row.cells[4]
            if entry.get("na_reason"):
                write_cell_paragraphs(
                    text_cell, [f"N/A — {entry['na_reason']}"]
                )
            else:
                section = entry.get("section_ref") or ""
                quote = entry.get("quote") or ""
                write_cell_paragraphs(
                    text_cell, [f"{section} — “{quote}”"]
                )
            filled_cells += 1
        table_stats.append({"matched": t_match, "skipped_variants": t_skip})

    doc.save(str(out_path))
    orphans = sorted(expected_keys - seen_keys)
    return {
        "matched": matched,
        "expected": len(expected_keys),
        "unmapped_rows": unmapped,
        "orphan_keys": orphans,
        "na_count": na_count,
        "filled_cells": filled_cells,
        "n_tables": len(doc.tables),
        "n_rows_by_table": [len(t.rows) for t in doc.tables],
        "per_table": table_stats,
    }


# ---------------------------------------------------------------------------
# TRIPOD filler
# ---------------------------------------------------------------------------


def fill_tripod_docx(map_dict: dict[str, dict], in_path: Path, out_path: Path) -> dict:
    """Fill TRIPOD Page (col 4)."""
    doc = Document(str(in_path))
    table = doc.tables[0]
    matched, unmapped, na_count, filled_cells = 0, [], 0, 0
    expected_keys = set(map_dict.keys())
    seen_keys: set[str] = set()

    for ri, row in enumerate(table.rows):
        if ri == 0:
            continue  # header row: Section/Topic | Item | (D/V blank) | Checklist Item | Page
        item_id = extract_item_id(row.cells[1].text)
        if item_id is None:
            continue  # Section band (Title and abstract / Introduction / Methods / …)
        entry = map_dict.get(item_id)
        if entry is None:
            unmapped.append((ri, row.cells[1].text.strip()))
            print(
                f"[TRIPOD] WARNING: row {ri} item '{item_id}' not in map",
                file=sys.stderr,
            )
            continue
        seen_keys.add(item_id)
        matched += 1
        dest = row.cells[4]
        if entry.get("na_reason"):
            write_cell_paragraphs(dest, [f"N/A — {entry['na_reason']}"])
            na_count += 1
        else:
            page = entry.get("page")
            section = entry.get("section_ref") or ""
            line = f"p. {page} — {section}".rstrip(" —")
            write_cell_paragraphs(dest, [line])
        filled_cells += 1

    doc.save(str(out_path))
    orphans = sorted(expected_keys - seen_keys)
    return {
        "matched": matched,
        "expected": len(expected_keys),
        "unmapped_rows": unmapped,
        "orphan_keys": orphans,
        "na_count": na_count,
        "filled_cells": filled_cells,
        "n_tables": len(doc.tables),
        "n_rows_by_table": [len(t.rows) for t in doc.tables],
    }


# ---------------------------------------------------------------------------
# Markdown companion writer
# ---------------------------------------------------------------------------


def _md_escape(s: str) -> str:
    if s is None:
        return ""
    return s.replace("|", "\\|").replace("\n", " ").replace("\r", " ").strip()


def _status_counts(map_dict: dict[str, dict]) -> tuple[int, int]:
    mapped = sum(1 for e in map_dict.values() if not e.get("na_reason"))
    na = sum(1 for e in map_dict.values() if e.get("na_reason"))
    return mapped, na


def _sort_items(keys: list[str]) -> list[str]:
    """Sort STROBE/CHEERS item ids: numeric part then letter suffix."""
    def k(x: str):
        m = re.match(r"(\d+)([a-z]?)", x)
        if not m:
            return (10**9, x)
        return (int(m.group(1)), m.group(2))
    return sorted(keys, key=k)


def write_markdown(
    map_dict: dict[str, dict],
    name: str,
    out_path: Path,
    columns: list[str],
) -> None:
    mapped, na = _status_counts(map_dict)
    lines: list[str] = []
    lines.append(f"# {name} — FORD-II v12")
    lines.append("")
    lines.append(f"**Manuscript:** {MANUSCRIPT_TITLE}")
    lines.append("")
    lines.append(f"**Generated:** {GEN_DATE}")
    lines.append("")
    lines.append(
        f"**Status counts:** Mapped = {mapped}; N/A = {na}; "
        f"Total = {mapped + na}"
    )
    lines.append("")
    lines.append("| " + " | ".join(columns) + " |")
    lines.append("| " + " | ".join(["---"] * len(columns)) + " |")

    has_dv = "D/V" in columns
    for key in _sort_items(list(map_dict.keys())):
        e = map_dict[key]
        item = e.get("item", key)
        topic = e.get("topic", "")
        rec = e.get("recommendation", "")
        page = e.get("page")
        section = e.get("section_ref") or ""
        if e.get("na_reason"):
            reported = f"N/A — {e['na_reason']}"
            page_str = "N/A"
            section_str = ""
        else:
            reported = e.get("quote") or ""
            page_str = f"p. {page}" if page is not None else ""
            section_str = section
        cols = [item]
        if has_dv:
            cols.append(e.get("applies_dv", ""))
        cols.extend([topic, rec, page_str, section_str, reported])
        lines.append("| " + " | ".join(_md_escape(c) for c in cols) + " |")

    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


# ---------------------------------------------------------------------------
# Output verification
# ---------------------------------------------------------------------------


def verify_outputs() -> dict[str, dict]:
    """Re-open each completed docx, assert tables/rows match templates, and
    confirm every destination cell is non-empty."""
    summary: dict[str, dict] = {}

    # CHEERS
    tpl = Document(str(CHEERS_TEMPLATE))
    out = Document(str(CHEERS_OUT_DOCX))
    assert len(tpl.tables) == len(out.tables) == 1, "CHEERS table count mismatch"
    assert len(tpl.tables[0].rows) == len(out.tables[0].rows), "CHEERS row count mismatch"
    cheers_filled = 0
    cheers_blank: list[int] = []
    for ri, row in enumerate(out.tables[0].rows):
        if ri == 0:
            continue
        item_id = extract_item_id(row.cells[1].text)
        if item_id is None:
            continue
        cell_txt = row.cells[3].text.strip()
        if cell_txt:
            cheers_filled += 1
        else:
            cheers_blank.append(ri)
    summary["cheers"] = {
        "tables_match": True,
        "rows_match": True,
        "filled": cheers_filled,
        "blank_rows": cheers_blank,
    }

    # STROBE
    tpl = Document(str(STROBE_TEMPLATE))
    out = Document(str(STROBE_OUT_DOCX))
    assert len(tpl.tables) == len(out.tables) == 3, "STROBE table count mismatch"
    for ti in range(3):
        assert len(tpl.tables[ti].rows) == len(out.tables[ti].rows), (
            f"STROBE table {ti} row count mismatch"
        )
    strobe_filled = 0
    strobe_blank: list[tuple[int, int]] = []
    for ti, table in enumerate(out.tables):
        for ri, row in enumerate(table.rows):
            item_no_text = row.cells[1].text.strip()
            if not any(ch.isdigit() for ch in item_no_text):
                continue
            rec_text = row.cells[2].text
            if item_no_text.startswith("15") and (
                "Case-control study" in rec_text or "Cross-sectional study" in rec_text
            ):
                continue  # skipped by design
            page_txt = row.cells[3].text.strip()
            text_txt = row.cells[4].text.strip()
            if page_txt and text_txt:
                strobe_filled += 1
            else:
                strobe_blank.append((ti, ri))
    summary["strobe"] = {
        "tables_match": True,
        "rows_match": True,
        "filled": strobe_filled,
        "blank_rows": strobe_blank,
    }

    # TRIPOD
    tpl = Document(str(TRIPOD_TEMPLATE))
    out = Document(str(TRIPOD_OUT_DOCX))
    assert len(tpl.tables) == len(out.tables) == 1, "TRIPOD table count mismatch"
    assert len(tpl.tables[0].rows) == len(out.tables[0].rows), "TRIPOD row count mismatch"
    tripod_filled = 0
    tripod_blank: list[int] = []
    for ri, row in enumerate(out.tables[0].rows):
        if ri == 0:
            continue
        item_id = extract_item_id(row.cells[1].text)
        if item_id is None:
            continue
        cell_txt = row.cells[4].text.strip()
        if cell_txt:
            tripod_filled += 1
        else:
            tripod_blank.append(ri)
    summary["tripod"] = {
        "tables_match": True,
        "rows_match": True,
        "filled": tripod_filled,
        "blank_rows": tripod_blank,
    }

    return summary


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> int:
    maps = load_maps()
    cheers_map = maps["cheers"]
    strobe_map = maps["strobe"]
    tripod_map = maps["tripod"]

    print("=" * 72)
    print("FORD-II v12 — Building reporting-guideline checklists")
    print("=" * 72)

    cheers_stats = fill_cheers_docx(cheers_map, CHEERS_TEMPLATE, CHEERS_OUT_DOCX)
    write_markdown(
        cheers_map,
        "CHEERS-2022 Checklist",
        CHEERS_OUT_MD,
        ["Item", "Topic", "Recommendation", "Page", "Section", "Reported text / N-A reason"],
    )

    strobe_stats = fill_strobe_docx(strobe_map, STROBE_TEMPLATE, STROBE_OUT_DOCX)
    write_markdown(
        strobe_map,
        "STROBE v4-combined (PLoS Medicine) Checklist",
        STROBE_OUT_MD,
        ["Item", "Topic", "Recommendation", "Page", "Section", "Reported text / N-A reason"],
    )

    tripod_stats = fill_tripod_docx(tripod_map, TRIPOD_TEMPLATE, TRIPOD_OUT_DOCX)
    write_markdown(
        tripod_map,
        "TRIPOD-2015 Checklist",
        TRIPOD_OUT_MD,
        [
            "Item",
            "D/V",
            "Topic",
            "Recommendation",
            "Page",
            "Section",
            "Reported text / N-A reason",
        ],
    )

    # Coverage report
    print()
    print("COVERAGE REPORT")
    print("-" * 72)
    print(
        f"CHEERS:  {cheers_stats['matched']} matched / "
        f"{cheers_stats['expected']} expected / "
        f"{len(cheers_stats['unmapped_rows'])} unmapped — "
        f"N/A={cheers_stats['na_count']}"
    )
    print(
        f"STROBE:  {strobe_stats['matched']} matched / "
        f"{strobe_stats['expected']} expected / "
        f"{len(strobe_stats['unmapped_rows'])} unmapped — "
        f"N/A={strobe_stats['na_count']}"
    )
    for ti, stats in enumerate(strobe_stats["per_table"]):
        print(
            f"   STROBE Table {ti}: matched={stats['matched']}, "
            f"skipped variant rows={stats['skipped_variants']}"
        )
    print(
        f"TRIPOD:  {tripod_stats['matched']} matched / "
        f"{tripod_stats['expected']} expected / "
        f"{len(tripod_stats['unmapped_rows'])} unmapped — "
        f"N/A={tripod_stats['na_count']}"
    )

    # Orphan items in the maps that never matched a template row
    print()
    print("ORPHAN MAP KEYS (no template row matched)")
    print("-" * 72)
    for name, stats in (
        ("CHEERS", cheers_stats),
        ("STROBE", strobe_stats),
        ("TRIPOD", tripod_stats),
    ):
        if stats["orphan_keys"]:
            print(f"  {name}: {stats['orphan_keys']}")
        else:
            print(f"  {name}: none")

    # Fail loudly if any cohort-branch row was unmapped
    any_unmapped = (
        bool(cheers_stats["unmapped_rows"])
        or bool(strobe_stats["unmapped_rows"])
        or bool(tripod_stats["unmapped_rows"])
    )
    if any_unmapped:
        print(
            "\nFATAL: at least one template row had no matching map entry — "
            "see warnings above.",
            file=sys.stderr,
        )
        return 1

    # Verification re-open
    print()
    print("VERIFICATION (re-open completed docx)")
    print("-" * 72)
    v = verify_outputs()
    for name in ("cheers", "strobe", "tripod"):
        info = v[name]
        print(
            f"  {name.upper():7s}: filled={info['filled']}, "
            f"blank={len(info['blank_rows'])}, "
            f"tables_match={info['tables_match']}, "
            f"rows_match={info['rows_match']}"
        )

    print()
    print("OUTPUT FILE SIZES")
    print("-" * 72)
    for p in (
        CHEERS_OUT_DOCX,
        CHEERS_OUT_MD,
        STROBE_OUT_DOCX,
        STROBE_OUT_MD,
        TRIPOD_OUT_DOCX,
        TRIPOD_OUT_MD,
    ):
        sz = p.stat().st_size if p.exists() else 0
        print(f"  {p.name:65s} {sz:>9d} B")

    print()
    print("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
