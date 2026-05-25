"""
Build TRIPOD-2015 22-item checklist supplement for the FORD-II manuscript.

Produces manuscript/ford_ii_tripod_supplement.docx — a 22-row table mapping
each TRIPOD-2015 item to its location in manuscript/ford_ii_main.docx, with
status (Addressed / Partially addressed / Not applicable).

TRIPOD-2015 reference:
    Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent Reporting of a
    multivariable prediction model for Individual Prognosis Or Diagnosis
    (TRIPOD): The TRIPOD Statement. Ann Intern Med. 2015;162:55-63.

Usage:
    python analysis/manuscript/06_build_tripod_supplement.py
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
HERE = Path(__file__).resolve().parent
REPO_ROOT = HERE.parents[1]
MANUSCRIPT_DIR = REPO_ROOT / "manuscript"
MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
MAIN_DOCX = MANUSCRIPT_DIR / "ford_ii_main.docx"
OUT_DOCX = MANUSCRIPT_DIR / "ford_ii_tripod_supplement.docx"

# Canonical FORD-II weights CSV (number of predictors derived from row count).
# Produced by analysis/ford_ii_refit/03_select_and_fit.py and shipped into
# REPO_ROOT/tables/ for distribution.
V3_WEIGHTS_CSV = (
    REPO_ROOT / "tables" / "ford_ii_weights.csv"
)


def _count_ford2_predictors() -> int:
    """Return the number of FORD-II predictors from the canonical weights CSV.

    Falls back to 16 if the CSV is not present (e.g., upstream pipeline not
    yet run when this script is invoked in isolation).
    """
    if not V3_WEIGHTS_CSV.exists():
        print(f"[WARN] {V3_WEIGHTS_CSV} missing; assuming 16 predictors.")
        return 16
    with V3_WEIGHTS_CSV.open() as f:
        # subtract 1 for header row
        return sum(1 for _ in f) - 1


N_FORD2_PREDICTORS = _count_ford2_predictors()


# ---------------------------------------------------------------------------
# TRIPOD-2015 22 items (published wording, lightly abbreviated where noted)
# ---------------------------------------------------------------------------
# Each entry: (item_id, title, description, default_section_ref, default_status)
# Section references resolved against ford_v4_jacs_v1.docx headings extracted at
# runtime (see resolve_sections() below).
TRIPOD_ITEMS = [
    # --- Title and abstract ---
    {
        "id": "1",
        "title": "Title",
        "description": (
            "Identify the study as developing and/or validating a multivariable "
            "prediction model, the target population, and the outcome to be predicted."
        ),
        "section": "Title page (separate file: ford_ii_titlepage.docx)",
        "status": "Addressed",
    },
    {
        "id": "2",
        "title": "Abstract",
        "description": (
            "Provide a summary of objectives, study design, setting, participants, "
            "sample size, predictors, outcome, statistical analysis, results, and "
            "conclusions."
        ),
        "section": "Abstract (Background / Methods / Results / Conclusions)",
        "status": "Addressed",
    },
    # --- Introduction ---
    {
        "id": "3a",
        "title": "Background and objectives",
        "description": (
            "Explain the medical context (including whether diagnostic or prognostic) "
            "and rationale for developing or validating the multivariable prediction "
            "model, including references to existing models."
        ),
        "section": "Introduction",
        "status": "Addressed",
    },
    {
        "id": "3b",
        "title": "Objectives",
        "description": (
            "Specify the objectives, including whether the study describes the "
            "development or validation of the model or both."
        ),
        "section": "Introduction (final paragraph) — external validation + cost-effectiveness of FORD-II",
        "status": "Addressed",
    },
    # --- Methods: Source of data ---
    {
        "id": "4a",
        "title": "Source of data",
        "description": (
            "Describe the study design or source of data (e.g., randomized trial, "
            "cohort, or registry data), separately for the development and validation "
            "data sets, if applicable."
        ),
        "section": "Methods § 4.1 Cohort and outcome (NTDB 2019–2024)",
        "status": "Addressed",
    },
    {
        "id": "4b",
        "title": "Key dates",
        "description": (
            "Specify the key study dates, including start of accrual, end of accrual, "
            "and, if applicable, end of follow-up."
        ),
        "section": "Methods § 4.1 (calendar years 2019–2024)",
        "status": "Addressed",
    },
    # --- Methods: Participants ---
    {
        "id": "5a",
        "title": "Participants — setting",
        "description": (
            "Specify key elements of the study setting (e.g., primary care, secondary "
            "care, general population), including number and location of centres."
        ),
        "section": "Methods § 4.1 (NTDB participating centres, U.S. trauma registry)",
        "status": "Addressed",
    },
    {
        "id": "5b",
        "title": "Participants — eligibility",
        "description": (
            "Describe eligibility criteria for participants."
        ),
        "section": "Methods § 4.1 (adult fracture inclusion; AMA/death/hospice/jail exclusions)",
        "status": "Addressed",
    },
    {
        "id": "5c",
        "title": "Participants — treatments received",
        "description": (
            "Give details of treatments received, if relevant."
        ),
        "section": "Not applicable — observational registry; no protocolised treatment exposure",
        "status": "Not applicable",
    },
    # --- Methods: Outcome ---
    {
        "id": "6a",
        "title": "Outcome definition",
        "description": (
            "Clearly define the outcome that is predicted by the prediction model, "
            "including how and when assessed."
        ),
        "section": "Methods § 4.1 (non-home discharge: REHAB/SNF/LTCH/ICF at index discharge)",
        "status": "Addressed",
    },
    {
        "id": "6b",
        "title": "Outcome blinding",
        "description": (
            "Report any actions to blind assessment of the outcome to be predicted."
        ),
        "section": (
            "Not applicable — outcome is administratively coded discharge disposition; "
            "no human adjudication"
        ),
        "status": "Not applicable",
    },
    # --- Methods: Predictors ---
    {
        "id": "7a",
        "title": "Predictors — definition",
        "description": (
            "Clearly define all predictors used in developing or validating the "
            "multivariable prediction model, including how and when they were measured."
        ),
        "section": (
            f"Methods § 4.2 Predictors and comparators ({N_FORD2_PREDICTORS} FORD-II predictors; "
            "weights from tables/ford_ii_weights.csv)"
        ),
        "status": "Addressed",
    },
    {
        "id": "7b",
        "title": "Predictors — blinding",
        "description": (
            "Report any actions to blind assessment of predictors for the outcome and "
            "other predictors."
        ),
        "section": (
            "Not applicable — predictors are pre-discharge administrative/clinical "
            "fields recorded prospectively into NTDB before disposition is finalised"
        ),
        "status": "Not applicable",
    },
    # --- Methods: Sample size ---
    {
        "id": "8",
        "title": "Sample size",
        "description": (
            "Explain how the study size was arrived at."
        ),
        "section": (
            "Methods § 4.1 (full NTDB 2019–2024 fracture cohort, N≈1.95M; no a priori "
            "power calculation — full-cohort validation)"
        ),
        "status": "Addressed",
    },
    # --- Methods: Missing data ---
    {
        "id": "9",
        "title": "Missing data",
        "description": (
            "Describe how missing data were handled (e.g., complete-case analysis, "
            "single imputation, multiple imputation) with details of any imputation "
            "method."
        ),
        "section": (
            "Methods § 4.2 — complete-case analysis on the FORD-II predictor set "
            "(LASSO refit in v3 used complete cases); see Limitations § 6.6"
        ),
        "status": "Partially addressed",
    },
    # --- Methods: Statistical analysis methods ---
    {
        "id": "10a",
        "title": "Statistical analysis — development continuous predictors",
        "description": (
            "(Development only) Describe how predictors were handled in the analyses."
        ),
        "section": (
            "Methods § 4.2 — FORD-II derivation pipeline lives in "
            "analysis/ford_ii_refit/; this manuscript applies and externally "
            "validates the locked model"
        ),
        "status": "Addressed",
    },
    {
        "id": "10b",
        "title": "Statistical analysis — model-building procedures",
        "description": (
            "(Development only) Specify type of model, all model-building procedures "
            "(including any predictor selection), and method for internal validation."
        ),
        "section": (
            "Methods § 4.2 — FORD-II is a LASSO refit (see analysis/ford_ii_refit/); "
            "this manuscript does not refit"
        ),
        "status": "Addressed",
    },
    {
        "id": "10c",
        "title": "Statistical analysis — validation methods",
        "description": (
            "(Validation only) For validation, describe how the predictions were "
            "calculated."
        ),
        "section": (
            "Methods § 4.3 (FORD-II linear predictor applied via locked weights from "
            "tables/ford_ii_weights.csv)"
        ),
        "status": "Addressed",
    },
    {
        "id": "10d",
        "title": "Statistical analysis — performance measures",
        "description": (
            "Specify all measures used to assess model performance and, if relevant, "
            "to compare multiple models."
        ),
        "section": (
            "Methods § 4.3 (AUROC + 1000-bootstrap 95% CI, DeLong test, Brier and "
            "scaled Brier, calibration slope/intercept, Hosmer–Lemeshow, decision-curve "
            "analysis)"
        ),
        "status": "Addressed",
    },
    {
        "id": "10e",
        "title": "Statistical analysis — model updating",
        "description": (
            "Describe any model updating (e.g., recalibration) arising from the "
            "validation, if done."
        ),
        "section": (
            "Methods § 4.3 — frozen FORD-II integer weights applied to the held-out "
            "test set without refit; no recalibration was applied"
        ),
        "status": "Addressed",
    },
    # --- Methods: Risk groups ---
    {
        "id": "11",
        "title": "Risk groups",
        "description": (
            "Provide details on how risk groups were created, if done."
        ),
        "section": (
            "Methods § 4.3 — risk quartiles defined on the FORD-II linear predictor"
        ),
        "status": "Addressed",
    },
    # --- Methods: Development vs validation ---
    {
        "id": "12",
        "title": "Development vs. validation",
        "description": (
            "(Validation only) For validation, identify any differences from the "
            "development data in setting, eligibility criteria, outcome, and "
            "predictors."
        ),
        "section": (
            "Methods § 4.4 Subgroup and sensitivity analyses; Discussion § 6.2 "
            "addresses comparison with prior literature"
        ),
        "status": "Addressed",
    },
    # --- Results: Participants ---
    {
        "id": "13a",
        "title": "Participants — flow",
        "description": (
            "Describe the flow of participants through the study, including the "
            "number of participants with and without the outcome and, if applicable, "
            "a summary of the follow-up time. A diagram may be helpful."
        ),
        "section": "Results § 5.1 Cohort and Figure 1 (CONSORT 7.33M → 1.95M)",
        "status": "Addressed",
    },
    {
        "id": "13b",
        "title": "Participants — characteristics",
        "description": (
            "Describe the characteristics of the participants (basic demographics, "
            "clinical features, available predictors), including the number of "
            "participants with missing data for predictors and outcome."
        ),
        "section": (
            "Results § 5.1 and Table 1 (baseline cohort characteristics, Train/Test)"
        ),
        "status": "Addressed",
    },
    {
        "id": "13c",
        "title": "Participants — comparison",
        "description": (
            "(Validation only) Show a comparison with the development data of the "
            "distribution of important variables (demographics, predictors, and "
            "outcome)."
        ),
        "section": (
            "Results § 5.1 / Table 1 — comparison across derivation and validation "
            "splits; Discussion § 6.2 compares with prior literature"
        ),
        "status": "Addressed",
    },
    # --- Results: Model development ---
    {
        "id": "14a",
        "title": "Model development — number of events",
        "description": (
            "(Development only) Specify the number of participants and outcome events "
            "in each analysis."
        ),
        "section": (
            "Methods § 4.1 / Results § 5.1 — cohort and event counts; full development "
            "details in analysis/ford_ii_refit/"
        ),
        "status": "Addressed",
    },
    {
        "id": "14b",
        "title": "Model development — unadjusted associations",
        "description": (
            "(Development only) If done, report the unadjusted association between "
            "each candidate predictor and outcome."
        ),
        "section": (
            "Not applicable in this validation manuscript — unadjusted associations "
            "are produced by analysis/ford_ii_refit/"
        ),
        "status": "Not applicable",
    },
    # --- Results: Model specification ---
    {
        "id": "15a",
        "title": "Model specification — full model",
        "description": (
            "(Development only) Present the full prediction model to allow predictions "
            "for individuals (i.e., all regression coefficients, and model intercept "
            "or baseline survival at a given time point)."
        ),
        "section": (
            "Results § 5.2 cites locked FORD-II weights table — "
            "tables/ford_ii_weights.csv"
        ),
        "status": "Addressed",
    },
    {
        "id": "15b",
        "title": "Model specification — explanation",
        "description": (
            "(Development only) Explain how to use the prediction model."
        ),
        "section": (
            "Results § 5.4 (risk quartile thresholds and event rates) and Discussion "
            "§ 6.3 Deployment implications (top-quartile workflow)"
        ),
        "status": "Addressed",
    },
    # --- Results: Model performance ---
    {
        "id": "16",
        "title": "Model performance",
        "description": (
            "Report performance measures (with CIs) for the prediction model."
        ),
        "section": (
            "Results § 5.2 (Table 2 discrimination + Figure 2 ROC) / § 5.3 (Table 3 "
            "calibration + Figure 3) / § 5.4 (Table 4 risk quartiles + Figure 4 DCA)"
        ),
        "status": "Addressed",
    },
    # --- Results: Model updating ---
    {
        "id": "17",
        "title": "Model updating",
        "description": (
            "(Validation only) If done, report the results from any model updating "
            "(i.e., model specification, model performance)."
        ),
        "section": "Not applicable — no model updating performed (FORD-II weights are locked)",
        "status": "Not applicable",
    },
    # --- Discussion ---
    {
        "id": "18",
        "title": "Limitations",
        "description": (
            "Discuss any limitations of the study (such as nonrepresentative sample, "
            "few events per predictor, missing data)."
        ),
        "section": (
            "Discussion § 6.5 Limitations (NTDB administrative coding, complete-case "
            "analysis, no causal LOS-reduction proof)"
        ),
        "status": "Addressed",
    },
    {
        "id": "19a",
        "title": "Interpretation — validation context",
        "description": (
            "(Validation only) For validation, discuss the results with reference to "
            "performance in the development data, and any other validation data."
        ),
        "section": (
            "Discussion § 6.1 Principal findings and § 6.2 Comparison with prior "
            "literature"
        ),
        "status": "Addressed",
    },
    {
        "id": "19b",
        "title": "Interpretation — overall",
        "description": (
            "Give an overall interpretation of the results, considering objectives, "
            "limitations, results from similar studies, and other relevant evidence."
        ),
        "section": (
            "Discussion § 6.1–6.2 (principal findings + comparison with prior "
            "literature) and § 6.4 Strengths"
        ),
        "status": "Addressed",
    },
    {
        "id": "20",
        "title": "Implications",
        "description": (
            "Discuss the potential clinical use of the model and implications for "
            "future research."
        ),
        "section": (
            "Discussion § 6.3 Deployment implications and Conclusions"
        ),
        "status": "Addressed",
    },
    # --- Other information ---
    {
        "id": "21",
        "title": "Supplementary information",
        "description": (
            "Provide information about the availability of supplementary resources, "
            "such as study protocol, web calculator, and data sets."
        ),
        "section": (
            "Supplements: ford_ii_supplement.docx; ford_ii_cheers_supplement.docx; "
            "this TRIPOD supplement; underlying CSVs in tables/"
        ),
        "status": "Addressed",
    },
    {
        "id": "22",
        "title": "Funding",
        "description": (
            "Give the source of funding and the role of the funders for the present "
            "study."
        ),
        "section": "Title page (separate file: ford_ii_titlepage.docx) — Funding statement",
        "status": "Addressed",
    },
]


# ---------------------------------------------------------------------------
# Cross-check: try to locate each "Methods § X.Y" / "Results § X.Y" /
# "Discussion § X.Y" anchor in the main manuscript so we can flag dead refs.
# ---------------------------------------------------------------------------
def collect_section_anchors(docx_path: Path) -> dict[str, int]:
    """Return {heading_text: paragraph_index} for short numbered headings.

    Captures top-level section headings (e.g., "Introduction", "Methods",
    "Results", "Discussion", "Conclusions", "Abstract", "References") and
    numbered subheadings ("4.1 Cohort and outcome", "5.2 Discrimination — head-
    to-head", etc.).
    """
    if not docx_path.exists():
        return {}
    doc = Document(str(docx_path))
    anchors: dict[str, int] = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if len(text) > 120:
            continue
        # Numbered subheadings: "4.1 ...", "5.2 ...", "6.3 ..."
        if len(text) >= 4 and text[0].isdigit() and text[1] == "." and text[2].isdigit():
            anchors[text] = i
            continue
        # Named top-level headings
        low = text.lower()
        if low in {
            "abstract",
            "introduction",
            "methods",
            "results",
            "discussion",
            "conclusions",
            "references",
        }:
            anchors[text] = i
    return anchors


def flag_unresolved(items: list[dict], anchors: dict[str, int]) -> list[str]:
    """Return list of item ids whose section reference does not match any
    anchor in the main manuscript (Title-page / N/A entries excluded)."""
    unresolved = []
    anchor_keys_lower = {k.lower(): k for k in anchors}
    for it in items:
        sec = it["section"].lower()
        if (
            it["status"] == "Not applicable"
            or "title page" in sec
            or "supplement" in sec
            or "abstract" in sec
        ):
            continue
        # Search for "4.1", "5.2", "6.3", etc., and "introduction"/"methods"/...
        hit = False
        for key_lower in anchor_keys_lower:
            if key_lower in sec:
                hit = True
                break
            # also try numeric prefix like "4.1"
            if key_lower[:3] in sec and key_lower[0].isdigit():
                hit = True
                break
        if not hit:
            unresolved.append(it["id"])
    return unresolved


# ---------------------------------------------------------------------------
# Word-table builder
# ---------------------------------------------------------------------------
def shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Calibri"


def status_color(status: str) -> str:
    return {
        "Addressed": "DCEAD2",          # light green
        "Partially addressed": "FFF2CC",  # light amber
        "Not applicable": "EAEAEA",       # neutral grey
    }.get(status, "FFFFFF")


def build_supplement(items: list[dict], unresolved: list[str], out_path: Path) -> int:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(
        "TRIPOD-2015 Checklist Supplement — FORD-II External Validation"
    )
    title_run.bold = True
    title_run.font.size = Pt(13)

    # Reference line
    ref = doc.add_paragraph()
    r1 = ref.add_run("Reference: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r2 = ref.add_run(
        "Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent Reporting of a "
        "multivariable prediction model for Individual Prognosis Or Diagnosis "
        "(TRIPOD): The TRIPOD Statement. Ann Intern Med. 2015;162:55-63."
    )
    r2.font.size = Pt(9)

    legend = doc.add_paragraph()
    r3 = legend.add_run("Status legend: ")
    r3.bold = True
    r3.font.size = Pt(9)
    r4 = legend.add_run(
        "Addressed = item fully reported in the cited section; "
        "Partially addressed = item reported with caveats (see section); "
        "Not applicable = item does not apply to an external-validation study with locked weights."
    )
    r4.font.size = Pt(9)

    # Build table: header + one row per item
    table = doc.add_table(rows=1 + len(items), cols=5)
    table.autofit = False
    table.style = "Table Grid"

    # Column widths
    col_widths = [Inches(0.5), Inches(1.6), Inches(2.6), Inches(2.4), Inches(1.0)]
    for col_idx, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[col_idx].width = width

    # Header row
    headers = ["Item #", "Item Title", "Description", "Section / Reference", "Status"]
    for c_idx, h in enumerate(headers):
        cell = table.rows[0].cells[c_idx]
        set_cell_text(cell, h, bold=True, size=9)
        shade_cell(cell, "B4C7E7")

    # Data rows
    for r_idx, it in enumerate(items, start=1):
        row = table.rows[r_idx]
        set_cell_text(row.cells[0], it["id"], bold=True, size=9)
        set_cell_text(row.cells[1], it["title"], bold=False, size=9)
        set_cell_text(row.cells[2], it["description"], bold=False, size=9)
        set_cell_text(row.cells[3], it["section"], bold=False, size=9)
        set_cell_text(row.cells[4], it["status"], bold=False, size=9)
        shade_cell(row.cells[4], status_color(it["status"]))

    # Footer: counts
    n_addressed = sum(1 for it in items if it["status"] == "Addressed")
    n_partial = sum(1 for it in items if it["status"] == "Partially addressed")
    n_na = sum(1 for it in items if it["status"] == "Not applicable")
    total = len(items)

    doc.add_paragraph()
    summary = doc.add_paragraph()
    sr = summary.add_run(
        f"Summary: {total} TRIPOD-2015 rows reported "
        f"(Addressed = {n_addressed}; Partially addressed = {n_partial}; "
        f"Not applicable = {n_na}). N/A items reflect that this study externally "
        "validates a locked-weight model on registry data without refit, recalibration, "
        "or model updating."
    )
    sr.font.size = Pt(9)
    sr.italic = True

    if unresolved:
        warn = doc.add_paragraph()
        wr = warn.add_run(
            "Unresolved section anchors (verify manually): "
            + ", ".join(unresolved)
        )
        wr.font.size = Pt(9)
        wr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        wr.italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return total


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    print(f"[INFO] Reading section anchors from: {MAIN_DOCX}")
    anchors = collect_section_anchors(MAIN_DOCX)
    print(f"[INFO] Found {len(anchors)} section anchors in main manuscript")

    unresolved = flag_unresolved(TRIPOD_ITEMS, anchors)
    if unresolved:
        print(f"[WARN] Could not auto-resolve section refs for items: {unresolved}")
    else:
        print("[INFO] All addressed/partial items map to a manuscript anchor")

    n = build_supplement(TRIPOD_ITEMS, unresolved, OUT_DOCX)

    n_addressed = sum(1 for it in TRIPOD_ITEMS if it["status"] == "Addressed")
    n_partial = sum(1 for it in TRIPOD_ITEMS if it["status"] == "Partially addressed")
    n_na = sum(1 for it in TRIPOD_ITEMS if it["status"] == "Not applicable")

    if not OUT_DOCX.exists():
        raise SystemExit(f"[FAIL] Output docx not written: {OUT_DOCX}")
    size_kb = OUT_DOCX.stat().st_size / 1024.0

    print()
    print("=" * 72)
    print(f"[OK] Wrote {OUT_DOCX} ({size_kb:.1f} KB)")
    print(f"[OK] Rows: {n} TRIPOD items")
    print(
        f"[OK] Status: Addressed={n_addressed}  Partial={n_partial}  N/A={n_na}"
    )
    na_ids = [it["id"] for it in TRIPOD_ITEMS if it["status"] == "Not applicable"]
    print(f"[OK] N/A item ids: {', '.join(na_ids)}")
    print("=" * 72)


if __name__ == "__main__":
    main()
