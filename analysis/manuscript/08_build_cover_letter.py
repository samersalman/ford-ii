"""
08_build_cover_letter.py

Build a templated cover-letter .docx for the FORD-II manuscript.

Design:
    - Letterhead-style layout, Times New Roman 11 pt, 1-inch margins,
      single-spaced with 6 pt paragraph spacing.
    - Square-bracket placeholders for any author-specific content
      (corresponding author name/title/email/phone, co-authors,
      affiliations, submission date, target journal).
    - Lead sentence + Maughan 2022 closest-precedent citation +
      why-this-journal justification + standard attestations
      (word count, COI, originality, prior publication).
    - Headline numerics (FORD-II AUROC 0.8285, n=650,737 test set,
      cohort N=1,952,210, Δcost cost-effectiveness summary) embedded
      verbatim.

Usage:
    python analysis/manuscript/08_build_cover_letter.py

Output:
    manuscript/ford_ii_cover_letter.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -------------------------------------------------------------------- paths
HERE = Path(__file__).resolve().parent
REPO_ROOT = HERE.parents[1]
MANUSCRIPT_DIR = REPO_ROOT / "manuscript"
MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = MANUSCRIPT_DIR / "ford_ii_cover_letter.docx"

# Canonical FORD-II weights CSV (number of predictors derived from row count).
# Produced by analysis/ford_ii_refit/03_select_and_fit.py.
V3_WEIGHTS_CSV = (
    REPO_ROOT
    / "tables"
    / "ford_ii_weights.csv"
)


def _count_ford2_predictors() -> int:
    """Return the number of FORD-II predictors from the canonical weights CSV.

    Falls back to 16 if the CSV is not present.
    """
    if not V3_WEIGHTS_CSV.exists():
        print(f"[WARN] {V3_WEIGHTS_CSV} missing; assuming 16 predictors.")
        return 16
    with V3_WEIGHTS_CSV.open() as f:
        # subtract 1 for header row
        return sum(1 for _ in f) - 1


N_FORD2_PREDICTORS = _count_ford2_predictors()

# -------------------------------------------------------------------- style
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(11)


def _style_doc(doc: Document) -> None:
    """1-inch margins; Times New Roman 11 pt; 6 pt spacing after."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)


def add_para(doc: Document, text: str, *, bold: bool = False,
             align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold


# -------------------------------------------------------------------- build
def build_cover_letter() -> Document:
    doc = Document()
    _style_doc(doc)

    # --- letterhead block (placeholders) ---
    add_para(doc, "[Corresponding Author Name], [Degree(s)]")
    add_para(doc, "[Title / Position]")
    add_para(doc, "[Department]")
    add_para(doc, "[Institution]")
    add_para(doc, "[Street Address]")
    add_para(doc, "[City, State ZIP]")
    add_para(doc, "Email: [Email]   |   Phone: [Phone]")
    add_para(doc, "")

    # --- date ---
    add_para(doc, "[Submission date]")
    add_para(doc, "")

    # --- recipient ---
    add_para(doc, "Editor-in-Chief")
    add_para(doc, "[Target Journal Name]")
    add_para(doc, "[Publisher / Society]")
    add_para(doc, "")

    # --- salutation ---
    add_para(doc, "Dear Editor-in-Chief,")
    add_para(doc, "")

    # --- lead paragraph ---
    add_para(
        doc,
        "We submit for your consideration the first decision-analytic "
        "cost-effectiveness analysis of a discharge-disposition prediction "
        "tool in adult trauma fractures. Our manuscript, "
        "“[Manuscript Title]”, presents the external validation "
        "and decision-analytic cost-effectiveness analysis of FORD-II, a "
        f"{N_FORD2_PREDICTORS}-predictor integer score for non-home discharge derived and "
        "validated on the National Trauma Data Bank (NTDB) 2019–2024 "
        "(analytic cohort N = 1,952,210; held-out test set n = 650,737; "
        "AUROC = 0.8285).",
    )

    # --- closest-precedent citation paragraph ---
    add_para(
        doc,
        "The closest methodological precedent is Maughan BC, et al., "
        "“Cost-Effectiveness of Field Trauma Triage among Injured Adults "
        "Served by Emergency Medical Services,” J Am Coll Surg "
        "2022;234:447–456 (PMID 35213435), a field-triage cost-"
        "effectiveness analysis. To our knowledge, no comparable decision-"
        "analytic evaluation has yet been published for an in-hospital "
        "discharge-disposition prediction tool in trauma surgery.",
    )

    # --- three-sentence why-this-journal justification ---
    add_para(
        doc,
        "We believe this work is an excellent fit for [Target Journal] for three reasons. "
        "First, FORD-II is a deployable surgical decision tool designed for "
        "bedside use on day-of-admission variables, directly addressing the "
        "discharge-planning bottleneck that surgical teams manage every day. "
        "Second, the manuscript adheres to dual TRIPOD (prediction-model) and "
        "CHEERS-2022 (cost-effectiveness) reporting standards, providing the "
        "methodological rigor expected of high-impact surgical submissions. "
        "Third, our decision-analytic model projects substantial annual "
        "savings if FORD-II were deployed nationally, a result of immediate "
        "policy and operational relevance.",
    )

    # --- key results paragraph ---
    add_para(
        doc,
        "Key headline results: FORD-II achieved an AUROC of 0.8285 on the "
        "650,737-row held-out NTDB test set; the hospital-perspective "
        "incremental cost favored the prediction-guided arm; 10,000-iteration "
        "probabilistic sensitivity analysis produced a high probability of "
        "net savings (see manuscript Results and Table 5).",
    )

    # --- standard attestations ---
    add_para(doc, "Standard attestations:", bold=True)

    add_para(
        doc,
        "Word count. The main text contains [insert word count] words "
        "(excluding abstract, references, tables, and figure legends), "
        "consistent with [Target Journal] Original Article guidelines.",
    )
    add_para(
        doc,
        "Originality and prior publication. This manuscript is original work. "
        "It has not been published elsewhere and is not under consideration "
        "by any other journal. The single-center derivation of the underlying "
        "FORD score was previously published in Injury (2026) and is cited "
        "and explicitly distinguished from the present FORD-II national "
        "external validation and cost-effectiveness analysis. No portion of "
        "the data, results, or analyses presented here has been previously "
        "published or submitted.",
    )
    add_para(
        doc,
        "Conflicts of interest. [The authors declare no relevant financial "
        "or non-financial conflicts of interest. / Disclosures: list here.]",
    )
    add_para(
        doc,
        "Authorship and contributions. All listed authors have read and "
        "approved the manuscript, meet ICMJE authorship criteria, and agree "
        "to be accountable for the integrity of the work. Author contributions "
        "are detailed on the title page.",
    )
    add_para(
        doc,
        "Funding. [Funding statement — e.g., “This work received no "
        "specific funding,” or list grant numbers.]",
    )
    add_para(
        doc,
        "Data and code availability. NTDB data are available from the "
        "American College of Surgeons under data-use agreement. Analysis code "
        "and the locked FORD-II score weights are available from the "
        "corresponding author upon reasonable request.",
    )
    add_para(
        doc,
        "IRB. This study used a de-identified, publicly available research "
        "dataset (NTDB 2019–2024) and was determined to be non-human-"
        "subjects research / exempt by the [Institution] Institutional Review "
        "Board (Protocol [IRB Number]).",
    )

    # --- co-authors / affiliations placeholders ---
    add_para(doc, "Authors and affiliations:", bold=True)
    add_para(
        doc,
        "Corresponding author: [Corresponding Author Name], [Degree(s)], "
        "[Title], [Institution], [Email], [Phone].",
    )
    add_para(
        doc,
        "Co-authors: [Co-Author 1, Degree(s)]; [Co-Author 2, Degree(s)]; "
        "[Co-Author 3, Degree(s)]; […].",
    )
    add_para(
        doc,
        "Affiliations: [Affiliation 1]; [Affiliation 2]; [Affiliation 3]; "
        "[…].",
    )

    # --- closing ---
    add_para(
        doc,
        "We hope you and the reviewers find this work suitable for "
        "publication in the Journal of the American College of Surgeons. "
        "Thank you for your time and consideration.",
    )
    add_para(doc, "")
    add_para(doc, "Sincerely,")
    add_para(doc, "")
    add_para(doc, "[Corresponding Author Name], [Degree(s)]")
    add_para(doc, "[Title / Position]")
    add_para(doc, "[Institution]")

    return doc


# -------------------------------------------------------------------- main
def _word_count(doc: Document) -> int:
    return sum(len(p.text.split()) for p in doc.paragraphs)


def _placeholder_count(doc: Document) -> int:
    """Count the [...] tokens left for the user to fill in."""
    import re
    pat = re.compile(r"\[[^\[\]]+\]")
    n = 0
    for p in doc.paragraphs:
        n += len(pat.findall(p.text))
    return n


if __name__ == "__main__":
    MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_cover_letter()
    doc.save(OUTPUT)

    assert OUTPUT.exists(), f"Failed to write {OUTPUT}"
    size_kb = OUTPUT.stat().st_size / 1024
    wc = _word_count(doc)
    ph = _placeholder_count(doc)

    print(f"Wrote: {OUTPUT}")
    print(f"  size       : {size_kb:.1f} KB")
    print(f"  word count : {wc}")
    print(f"  placeholders: {ph}")
    assert size_kb > 5, f"Output too small ({size_kb:.1f} KB <= 5 KB)"
    print("OK")
