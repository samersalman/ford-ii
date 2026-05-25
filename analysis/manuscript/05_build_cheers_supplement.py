"""
05_build_cheers_supplement.py — CHEERS-2022 reporting checklist supplement
for the FORD-II manuscript.

Builds a 28-row table mapping every CHEERS 2022 item to:
  - the section / paragraph reference in the main manuscript .docx
  - status (Addressed / Not applicable / Deferred)

Page numbers cannot be read from a python-docx document object (Word does not
store rendered pagination), so references use the manuscript's named
subsection (e.g., "Methods § 4.5 Cost-effectiveness model").

Inputs (read-only):
    figures/cost_analysis/cheers_2022_checklist.md
    manuscript/ford_ii_main.docx

Output:
    manuscript/ford_ii_cheers_supplement.docx

Notes:
  - Item 15 (Currency, price date, conversion) carries a verbatim Arm F
    PHCE non-material disclosure note.
  - Items 27, 28 (Funding, COI) are Deferred with a User-action note.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parents[1]
COST_DIR = REPO_ROOT / "figures" / "cost_analysis"
MANUSCRIPT_DIR = REPO_ROOT / "manuscript"
MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)

CHEERS_MD = COST_DIR / "cheers_2022_checklist.md"
MAIN_DOCX = MANUSCRIPT_DIR / "ford_ii_main.docx"
OUT_DOCX = MANUSCRIPT_DIR / "ford_ii_cheers_supplement.docx"


# ---------------------------------------------------------------------------
# Verbatim disclosures from the Step 5a plan
# ---------------------------------------------------------------------------

ARM_F_PHCE_NOTE = (
    "An alternative inflator using the CMS NHEA Personal Health Care "
    "Expenditure price index (PHCE 2021—2024 = 1.0821) was assessed in "
    "a one-way sensitivity analysis (Arm F); the resulting per-patient Δ"
    "cost ($366.19) differed from the base case ($362.25) by 1.09%, "
    "confirming that the choice of medical-cost inflator is non-material "
    "over the 2021—2024 horizon."
)

COI_DEFERRED_NOTE = (
    "User to populate Title-page funding/COI section before submission."
)


# ---------------------------------------------------------------------------
# CHEERS markdown parser
# ---------------------------------------------------------------------------

def parse_cheers_md(md_path: Path) -> list[dict]:
    """Return a list of 28 dicts: {item, section, name, description}.

    The checklist is a markdown table with header row plus 28 item rows.
    """
    text = md_path.read_text(encoding="utf-8")
    items = []
    for line in text.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        # skip header/separator rows
        if "Item #" in line or set(line.replace("|", "").strip()) <= set("- :"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) < 6:
            continue
        # Cells: 0 item#, 1 Section, 2 Item name, 3 Description, 4 Status, 5 Cross-ref
        try:
            num = int(cells[0])
        except ValueError:
            continue
        items.append(
            {
                "item": num,
                "section": cells[1],
                "name": cells[2],
                "description": cells[3],
                "status_md": cells[4],
                "xref_md": cells[5],
            }
        )
    if len(items) != 28:
        raise RuntimeError(
            f"Expected 28 CHEERS items, parsed {len(items)} from {md_path}"
        )
    return items


# ---------------------------------------------------------------------------
# Manuscript section map
# ---------------------------------------------------------------------------
# References use named subsections of the main manuscript .docx. The main
# manuscript builder produces these section headings; the CHEERS supplement
# does not require paragraph-index parity with the main doc.

MANUSCRIPT_SECTIONS = {
    "title":       "Title page",
    "abstract":    "Structured Abstract",
    "intro":       "Introduction",
    "methods":     "Methods",
    "m_4_1":       "Methods § 4.1 Cohort and outcome",
    "m_4_2":       "Methods § 4.2 Predictors and comparators",
    "m_4_3":       "Methods § 4.3 Statistical analysis",
    "m_4_4":       "Methods § 4.4 Subgroup and sensitivity analyses",
    "m_4_5":       "Methods § 4.5 Cost-effectiveness model (CHEERS 2022)",
    "m_4_6":       "Methods § 4.5 Cost-effectiveness model (CHEERS 2022)",
    "results":     "Results",
    "r_5_1":       "Results § 5.1 Cohort",
    "r_5_7":       "Results § 5.6 Cost analysis",
    "r_5_8":       "Results § 5.6 Cost analysis",
    "discussion":  "Discussion",
    "d_6_1":       "Discussion § 6.1 Principal findings",
    "d_6_2":       "Discussion § 6.2 Comparison with prior literature",
    "d_6_4":       "Discussion § 6.3 Deployment implications",
    "d_6_5":       "Discussion § 6.4 Strengths",
    "d_6_6":       "Discussion § 6.5 Limitations",
    "conclusions": "Conclusions",
    "tables":      "Tables",
    "table5":      "Table 5 (Cost-effectiveness summary)",
    "figures":     "Figures",
    "figure5":     "Figure 5 (Cost-effectiveness panel)",
    "titlepage_funding": "Title page (separate file: ford_ii_titlepage.docx)",
}


def sec(key: str) -> str:
    """Return the human-readable section reference for a manuscript anchor key."""
    if key not in MANUSCRIPT_SECTIONS:
        raise KeyError(f"Unknown manuscript section key: {key}")
    return MANUSCRIPT_SECTIONS[key]


# ---------------------------------------------------------------------------
# Per-item resolution map (status + section reference + note)
# ---------------------------------------------------------------------------
# Items 1-3, 26-28: Session-7 dependencies resolved against the manuscript.
# Item 15: includes the verbatim Arm F PHCE disclosure.
# Items 4-14, 16-25: section/page mapped per CHEERS section against Methods/Results.

ITEM_RESOLUTION: dict[int, dict] = {
    # Title — resolves against the manuscript title text on page 1
    1: {
        "status": "Addressed",
        "ref": (
            sec("title")
            + ' — "FORD-II: External Validation and Cost-Effectiveness '
            'of a Fracture Trauma Discharge-Disposition Score on the National '
            'Trauma Data Bank, 2019—2024" identifies the study as an '
            "external validation plus economic evaluation."
        ),
    },
    # Abstract — structured abstract present per JACS template
    2: {
        "status": "Addressed",
        "ref": (
            sec("abstract")
            + " — Structured (Background / Methods / Results / "
            "Conclusions) on page 1; covers context, key methods, base-case "
            "and PSA results, and sensitivity arms."
        ),
    },
    # Background and objectives — Introduction
    3: {
        "status": "Addressed",
        "ref": (
            sec("intro")
            + " — covers discharge-planning burden in trauma, GTOS-II/TRIAGES "
            "limitations, and the FORD-II national refit with an explicit "
            "decision-making/policy hook."
        ),
    },
    # Health economic analysis plan
    4: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + "; protocol locked in cost_analysis/methodology_lock.md §1."
        ),
    },
    # Study population
    5: {
        "status": "Addressed",
        "ref": (
            sec("m_4_1")
            + " (NTDB 2019—2024 adult fracture cohort; N=1,952,210); "
            + sec("r_5_1") + "; Table 1 baseline."
        ),
    },
    # Setting and location
    6: {
        "status": "Addressed",
        "ref": (
            sec("m_4_1")
            + " (NTDB / TQIP 2019—2024, U.S. trauma centers); "
            + sec("m_4_6") + " (hospital perspective; national rollout scenario)."
        ),
    },
    # Comparators
    7: {
        "status": "Addressed",
        "ref": (
            sec("m_4_2")
            + " (FORD-II vs. GTOS-II vs. TRIAGES); "
            + sec("m_4_6")
            + " (decision-tree two-arm comparator: top-quartile flag vs. "
            "usual care)."
        ),
    },
    # Perspective
    8: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + " (hospital perspective primary; payer perspective secondary)."
        ),
    },
    # Time horizon
    9: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + " (index admission + 30-day horizon; rationale in "
            "methodology_lock.md §3)."
        ),
    },
    # Discount rate
    10: {
        "status": "Not applicable",
        "ref": (
            sec("m_4_6")
            + " — horizon ≤1 year, so no discounting required "
            "(per CHEERS 2022 guidance)."
        ),
    },
    # Selection of outcomes
    11: {
        "status": "Addressed",
        "ref": (
            sec("m_4_1")
            + " (non-home discharge as the disposition outcome); "
            + sec("m_4_6")
            + " (cost Δ and length-of-stay reduction as economic "
            "outcomes)."
        ),
    },
    # Measurement of outcomes
    12: {
        "status": "Addressed",
        "ref": (
            sec("m_4_1")
            + " (NTDB DISCHARGEDISPOSITION coding; non-home = REHAB / SNF / "
            "LTCH / ICF); methodology_lock.md §4."
        ),
    },
    # Valuation of outcomes
    13: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + " (binary disposition; cost Δ valued in 2024 USD); "
            "methodology_lock.md §4."
        ),
    },
    # Measurement and valuation of resources and costs
    14: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + "; cost_model_inputs.csv (10-row parameter table); "
            "table_base_case.csv."
        ),
    },
    # Currency, price date, and conversion — INCLUDES verbatim Arm F PHCE
    15: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + " (2024 USD; BLS CPI-MED inflator 2021—2024 = 1.0710); "
            + sec("r_5_8")
            + " and Supplemental Table S6 Arm F. "
            + ARM_F_PHCE_NOTE
        ),
    },
    # Rationale and description of model
    16: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + " (two-arm decision-tree; FORD-II top-quartile flag triggers "
            "structured early discharge planning); cost_analysis/"
            "decision_tree_diagram.png; methodology_lock.md §6."
        ),
    },
    # Analytics and assumptions
    17: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + " (10,000-iter PSA, Briggs TF-6, seed 42, Normal LOS truncated "
            "at 0); methodology_lock.md §7."
        ),
    },
    # Characterizing heterogeneity
    18: {
        "status": "Addressed",
        "ref": (
            sec("m_4_4")
            + " (subgroup AUROCs); "
            + sec("r_5_8")
            + " (sensitivity arms by parameter); table_scaling.csv "
            "(per-center / national rollout)."
        ),
    },
    # Distributional effects
    19: {
        "status": "Not applicable",
        "ref": (
            "No distributional or equity-weighted CEA performed in FORD-II "
            "Session 6; documented in methodology_lock.md §3."
        ),
    },
    # Characterizing uncertainty
    20: {
        "status": "Addressed",
        "ref": (
            sec("m_4_6")
            + " and "
            + sec("r_5_7")
            + "; tornado (one-way SA), 10K-iter PSA, threshold analysis, CEAC; "
            "table_tornado.csv, table_threshold.csv, psa_results.csv; "
            "Figure 5 (panels A and B); Supplemental Figure S-Fig 4 (CEAC)."
        ),
    },
    # Patient/public engagement
    21: {
        "status": "Not applicable",
        "ref": (
            "No patient or public engagement conducted; this is a secondary-"
            "data cost model on a de-identified national registry (NTDB "
            "2019—2024)."
        ),
    },
    # Study parameters
    22: {
        "status": "Addressed",
        "ref": (
            sec("r_5_7")
            + " and Table 5; cost_model_inputs.csv (parameter values, ranges, "
            "distributions); psa_results.csv (10,000 PSA draws)."
        ),
    },
    # Summary of main results
    23: {
        "status": "Addressed",
        "ref": (
            sec("r_5_7")
            + " and Table 5 — deterministic Δcost $362.25/patient; "
            "PSA mean $361.63 (95% CrI −$6.53 to $894.62); "
            "P(net savings>0)=96.82%; cost-neutrality threshold 0.0468 d."
        ),
    },
    # Effect of uncertainty
    24: {
        "status": "Addressed",
        "ref": (
            sec("r_5_8")
            + " and Supplemental Table S6 (six pre-specified sensitivity "
            "arms incl. Arm F PHCE inflator); Figure 5A tornado; Figure 5B "
            "PSA scatter; Supplemental Figure S-Fig 4 CEAC."
        ),
    },
    # Effect of patient/public engagement
    25: {
        "status": "Not applicable",
        "ref": "See Item 21 — no patient or public engagement conducted.",
    },
    # Discussion of findings
    26: {
        "status": "Addressed",
        "ref": (
            sec("d_6_1") + " (principal findings); "
            + sec("d_6_2") + " (prior literature, Maughan 2022 *JACS* CEA "
            "precedent); "
            + sec("d_6_4") + " (deployment implications); "
            + sec("d_6_5") + " (strengths); "
            + sec("d_6_6") + " (limitations, ethical/equity considerations)."
        ),
    },
    # Source of funding
    27: {
        "status": "Deferred",
        "ref": (
            sec("titlepage_funding")
            + " — funding statement placeholder pending user input. "
            + COI_DEFERRED_NOTE
        ),
    },
    # Conflicts of interest
    28: {
        "status": "Deferred",
        "ref": (
            sec("titlepage_funding")
            + " — COI disclosures placeholder pending user input. "
            + COI_DEFERRED_NOTE
        ),
    },
}


# ---------------------------------------------------------------------------
# Docx helpers (port from FORD/v7/build_v7_manuscript.py patterns)
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, fill_hex: str) -> None:
    """Apply background shading to a table cell (light gray for header)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, size_pt: int = 9) -> None:
    """Replace cell paragraph(s) with a single paragraph carrying `text`."""
    cell.text = ""  # clears default empty paragraph
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _set_table_borders(table) -> None:
    """Apply single-line borders to all cells of a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "808080")
        tblBorders.append(b)
    # Replace any existing borders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_supplement() -> dict:
    if not CHEERS_MD.is_file():
        raise FileNotFoundError(f"CHEERS source missing: {CHEERS_MD}")
    if not MAIN_DOCX.is_file():
        raise FileNotFoundError(f"Main manuscript missing: {MAIN_DOCX}")

    # Sanity-check the main manuscript exists and is FORD-II
    main = Document(str(MAIN_DOCX))
    n_paras = len(main.paragraphs)
    if n_paras < 40:
        raise RuntimeError(
            f"Main manuscript has only {n_paras} paragraphs; expected ≥40. "
            "Re-run 04_build_manuscript.py."
        )
    title_text = main.paragraphs[0].text.lower()
    if "ford-ii" not in title_text:
        raise RuntimeError(
            f"Manuscript paragraph 0 does not contain 'FORD-II'; "
            f"got: {title_text[:80]!r}."
        )

    # Parse the CHEERS markdown
    items = parse_cheers_md(CHEERS_MD)

    # Build the supplement document
    doc = Document()

    # Page setup: letter, 1-in margins, landscape for the wide table
    section = doc.sections[0]
    section.page_height = Inches(8.5)
    section.page_width = Inches(11.0)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Header / preamble
    h = doc.add_heading("Supplemental Material — CHEERS 2022 Reporting Checklist", level=1)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "FORD-II: External Validation and Cost-Effectiveness of a Fracture "
        "Trauma Discharge-Disposition Score on the National Trauma Data Bank, "
        "2019—2024."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    intro = doc.add_paragraph()
    intro.add_run(
        "Mapping of every Consolidated Health Economic Evaluation Reporting "
        "Standards (CHEERS 2022) checklist item to its location in the main "
        "manuscript (ford_ii_main.docx). Section references use the "
        "manuscript's numbered subsection scheme (e.g., Methods § 4.5). "
        "Word does not store rendered page numbers, so named section "
        "references rather than page numbers are provided. Source: "
        "Husereau D, Drummond M, Augustovski F, et al. CHEERS 2022 "
        "Statement. BMJ 2022;376:e067975 (PMID 35017145; "
        "doi:10.1136/bmj-2021-067975)."
    ).font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Build the 28-item table (1 header + 28 rows)
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    _set_table_borders(table)

    # Column widths (sum = 9.5 in landscape inside 11in - 1.5 in margins)
    col_widths = [Inches(0.6), Inches(2.6), Inches(5.0), Inches(1.3)]

    hdr = table.rows[0].cells
    headers = ["Item #", "CHEERS Item Description", "Section / Reference", "Status"]
    for i, (cell, txt, w) in enumerate(zip(hdr, headers, col_widths)):
        _set_cell_text(cell, txt, bold=True, size_pt=10)
        _set_cell_shading(cell, "D9D9D9")
        cell.width = w

    # Counts for the report
    counts = {"Addressed": 0, "Not applicable": 0, "Deferred": 0}
    deferred_items = []
    placeholder_items = []
    unresolved = []

    for it in items:
        num = it["item"]
        resolution = ITEM_RESOLUTION.get(num)
        if resolution is None:
            unresolved.append(num)
            status = "Deferred"
            ref = "Section reference not resolved — user review required."
        else:
            status = resolution["status"]
            ref = resolution["ref"]

        counts[status] = counts.get(status, 0) + 1
        if status == "Deferred":
            deferred_items.append(num)
            if "User to populate" in ref or "pending user input" in ref.lower():
                placeholder_items.append(num)

        row = table.add_row().cells
        # Reset widths on the new row (Word loses them otherwise)
        for cell, w in zip(row, col_widths):
            cell.width = w

        # Compose the description: "Section: Item name — description"
        desc = (
            f"{it['section']}: {it['name']} — {it['description']}"
        )
        _set_cell_text(row[0], str(num), bold=True, size_pt=9)
        _set_cell_text(row[1], desc, bold=False, size_pt=9)
        _set_cell_text(row[2], ref, bold=False, size_pt=9)
        _set_cell_text(row[3], status, bold=(status == "Deferred"), size_pt=9)
        # Light yellow shading for Deferred rows so user-action items pop
        if status == "Deferred":
            _set_cell_shading(row[3], "FFF2CC")

    # Footer summary
    doc.add_paragraph()
    summary = doc.add_paragraph()
    summary_run = summary.add_run("Summary. ")
    summary_run.bold = True
    summary_run.font.size = Pt(10)
    summary.add_run(
        f"{counts['Addressed']} items Addressed; "
        f"{counts['Not applicable']} Not applicable; "
        f"{counts['Deferred']} Deferred. "
    ).font.size = Pt(10)
    if placeholder_items:
        ph_str = ", ".join(str(n) for n in placeholder_items)
        summary.add_run(
            f"User-action required to finalize Items {ph_str} (Title-page "
            "funding/COI placeholders pending input)."
        ).font.size = Pt(10)

    doc.save(str(OUT_DOCX))

    # Compute table row count post-write for the report
    written = Document(str(OUT_DOCX))
    written_table = written.tables[0]
    n_rows = len(written_table.rows)

    return {
        "items_total": len(items),
        "table_rows": n_rows,
        "counts": counts,
        "deferred_items": deferred_items,
        "placeholder_items": placeholder_items,
        "unresolved": unresolved,
        "out_size_kb": OUT_DOCX.stat().st_size / 1024,
    }


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    info = build_supplement()
    print(f"Wrote {OUT_DOCX}")
    print(f"  size: {info['out_size_kb']:,.1f} KB")
    print(f"  table rows: {info['table_rows']} (expected 29 = 1 header + 28 items)")
    print(f"  CHEERS items parsed: {info['items_total']}")
    print(f"  status counts: {info['counts']}")
    if info["placeholder_items"]:
        print(
            "  user-action items: "
            + ", ".join(str(n) for n in info["placeholder_items"])
        )
    if info["unresolved"]:
        print(
            "  unresolved items (need user review): "
            + ", ".join(str(n) for n in info["unresolved"])
        )
    if info["table_rows"] != 29:
        raise SystemExit(
            f"FAIL: table has {info['table_rows']} rows; expected 29."
        )
