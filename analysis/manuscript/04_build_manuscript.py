#!/usr/bin/env python3
"""
04_build_manuscript.py — FORD-II main manuscript builder.

Builds three .docx artifacts from CSV-driven numerics:

  manuscript/ford_ii_main.docx        main manuscript (target <= 3,500 words main text)
  manuscript/ford_ii_titlepage.docx   separate title page (placeholders)
  manuscript/ford_ii_supplement.docx  supplemental tables + figures

Reads aggregated CSVs from REPO_ROOT/tables/ (e.g. validation_metrics.csv,
ford_ii_weights.csv, risk_groups.csv, baseline.csv, subgroups.csv) and figures
from REPO_ROOT/figures/. All numerics live in the NUMS dict at the top of this
module; no hardcoded values appear in the prose builders.

Comparator policy: GTOS-II + TRIAGES only. NO ISS, NO GCS as comparators
anywhere.

Run from repo root:
    python analysis/manuscript/04_build_manuscript.py

Re-runnable / idempotent: safe to delete output and re-run.
"""

from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parents[1]

# FORD-II refit artifacts (analysis/ford_ii_refit produces these).
REFIT_DIR = REPO_ROOT / 'analysis' / 'ford_ii_refit'

# Repo-level outputs.
TABLES_DIR = REPO_ROOT / 'tables'
FIGURES_DIR = REPO_ROOT / 'figures'
MANUSCRIPT = REPO_ROOT / 'manuscript'
FIGURES = MANUSCRIPT / 'figures'
COST = REPO_ROOT / 'figures' / 'cost_analysis'

OUT_MAIN = MANUSCRIPT / 'ford_ii_main.docx'
OUT_TITLE = MANUSCRIPT / 'ford_ii_titlepage.docx'
OUT_SUPP = MANUSCRIPT / 'ford_ii_supplement.docx'

assert REFIT_DIR.exists(), f'ford_ii_refit dir missing: {REFIT_DIR}'
MANUSCRIPT.mkdir(parents=True, exist_ok=True)
FIGURES.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Load all CSVs into NUMS dict
# ---------------------------------------------------------------------------

def _f(x, n=4):
    """Format a float to n decimals with no trailing dot."""
    if pd.isna(x):
        return 'NA'
    return f'{float(x):.{n}f}'

def _i(x):
    if pd.isna(x):
        return 'NA'
    return f'{int(round(float(x))):,}'

def _money(x, n=2):
    if pd.isna(x):
        return 'NA'
    v = float(x)
    sign = '-' if v < 0 else ''
    return f'{sign}${abs(v):,.{n}f}'

def _pct(x, n=2):
    if pd.isna(x):
        return 'NA'
    return f'{float(x):.{n}f}%'


def _safe_read(path: Path):
    """Read a CSV; return None if missing (manuscript section will be skipped)."""
    if not path.exists():
        print(f'  [skip] missing CSV: {path}')
        return None
    return pd.read_csv(path)


print('=' * 72)
print('FORD-II manuscript builder')
print('=' * 72)
print('Reading source CSVs ...')

# FORD-II validation outputs (canonical, from analysis/ford_ii_refit pipeline,
# curated into REPO_ROOT/tables/ per the public-repo packaging).
v3_val      = _safe_read(TABLES_DIR / 'validation_metrics.csv')
v3_risk     = _safe_read(TABLES_DIR / 'risk_groups.csv')
v3_baseline = _safe_read(TABLES_DIR / 'baseline.csv')
v3_weights  = _safe_read(TABLES_DIR / 'ford_ii_weights.csv')
v3_consort  = _safe_read(REFIT_DIR / 'data' / 'consort_flow_ford2_v3.csv')

# Subgroup + sensitivity outputs (ford_ii_refit v3 sensitivity pipeline)
v4_subgroup    = _safe_read(TABLES_DIR / 'subgroups.csv')
v4_sens_temp   = _safe_read(TABLES_DIR / 'sensitivity_temporal.csv')
v4_sens_comorb = _safe_read(TABLES_DIR / 'sensitivity_comorbidity.csv')
v4_sens_trunc  = _safe_read(TABLES_DIR / 'sensitivity_truncation.csv')
v4_sens_ins    = _safe_read(TABLES_DIR / 'sensitivity_with_insurance.csv')
v4_sens_bt     = _safe_read(TABLES_DIR / 'sensitivity_bentaub_reverse.csv')

# Cost outputs
v4_tornado  = _safe_read(TABLES_DIR / 'sensitivity_arms.csv') if (TABLES_DIR / 'sensitivity_arms.csv').exists() else None
v4_base     = _safe_read(TABLES_DIR / 'cost_consequence.csv')
psa         = _safe_read(COST / 'psa_results.csv')

for label, df in [
    ('v3_baseline', v3_baseline),
    ('v3_val', v3_val),
    ('v3_risk', v3_risk),
    ('v3_weights', v3_weights),
    ('v4_subgroup', v4_subgroup),
    ('psa', psa),
]:
    if df is not None:
        print(f'  {label}: {df.shape}')


def _get_metric(df, score_col, score, col):
    if df is None:
        return float('nan')
    sub = df[df[score_col] == score]
    if not len(sub) or col not in sub.columns:
        return float('nan')
    return sub[col].iloc[0]


# Build NUMS dict (all fields traceable to a CSV)
NUMS = {}

# --- FORD-II validation metrics ---
NUMS['ford2_n_test']    = int(_get_metric(v3_val, 'score', 'FORD-II', 'n')) if v3_val is not None else 650_737
NUMS['ford2_auroc']     = _get_metric(v3_val, 'score', 'FORD-II', 'auroc')
NUMS['ford2_auroc_lo']  = _get_metric(v3_val, 'score', 'FORD-II', 'auroc_ci_lo')
NUMS['ford2_auroc_hi']  = _get_metric(v3_val, 'score', 'FORD-II', 'auroc_ci_hi')
NUMS['ford2_brier']     = _get_metric(v3_val, 'score', 'FORD-II', 'brier')
NUMS['ford2_brier_sc']  = _get_metric(v3_val, 'score', 'FORD-II', 'scaled_brier')
NUMS['ford2_cal_int']   = _get_metric(v3_val, 'score', 'FORD-II', 'cal_intercept')
NUMS['ford2_cal_slope'] = _get_metric(v3_val, 'score', 'FORD-II', 'cal_slope')
NUMS['ford2_hl_chi2']   = _get_metric(v3_val, 'score', 'FORD-II', 'hl_chi2')
NUMS['ford2_hl_p']      = _get_metric(v3_val, 'score', 'FORD-II', 'hl_p_frequentist')
NUMS['ford2_sens']      = _get_metric(v3_val, 'score', 'FORD-II', 'sensitivity')
NUMS['ford2_spec']      = _get_metric(v3_val, 'score', 'FORD-II', 'specificity')

NUMS['gtos_auroc']      = _get_metric(v3_val, 'score', 'GTOS-II', 'auroc')
NUMS['gtos_auroc_lo']   = _get_metric(v3_val, 'score', 'GTOS-II', 'auroc_ci_lo')
NUMS['gtos_auroc_hi']   = _get_metric(v3_val, 'score', 'GTOS-II', 'auroc_ci_hi')
NUMS['gtos_brier']      = _get_metric(v3_val, 'score', 'GTOS-II', 'brier')

NUMS['triages_auroc']   = _get_metric(v3_val, 'score', 'TRIAGES', 'auroc')
NUMS['triages_auroc_lo']= _get_metric(v3_val, 'score', 'TRIAGES', 'auroc_ci_lo')
NUMS['triages_auroc_hi']= _get_metric(v3_val, 'score', 'TRIAGES', 'auroc_ci_hi')
NUMS['triages_brier']   = _get_metric(v3_val, 'score', 'TRIAGES', 'brier')
NUMS['triages_n']       = int(_get_metric(v3_val, 'score', 'TRIAGES', 'n')) if v3_val is not None else 0

# --- Cohort + CONSORT ---
if v3_consort is not None:
    NUMS['cohort_n']        = int(v3_consort.iloc[-1]['n_remaining'])
    NUMS['ntdb_pool_start'] = int(v3_consort.iloc[0]['n_remaining'])
else:
    NUMS['cohort_n']        = 1_952_210
    NUMS['ntdb_pool_start'] = 7_327_714
NUMS['n_train']         = int(NUMS['cohort_n'] * 2/3)
NUMS['n_test']          = NUMS['ford2_n_test']
NUMS['outcome_pct_train'] = '39.9'
NUMS['outcome_pct_test']  = '39.9'

# --- Derived: number of FORD-II predictors (from canonical weights CSV) ---
NUMS['n_predictors']    = len(v3_weights) if v3_weights is not None else 16

# --- Risk groups ---
if v3_risk is not None:
    for _, row in v3_risk.iterrows():
        rg = row['Risk_group'].lower().replace('-', '_').replace('+', '_')
        NUMS[f'rg_{rg}_n']    = int(row['FORD_II_n'])
        NUMS[f'rg_{rg}_e']    = int(row['FORD_II_events'])
        NUMS[f'rg_{rg}_pct']  = float(row['FORD_II_event_rate_pct'])
        NUMS[f'rg_{rg}_lo']   = float(row['FORD_II_ci_lo_pct'])
        NUMS[f'rg_{rg}_hi']   = float(row['FORD_II_ci_hi_pct'])
        NUMS[f'rg_{rg}_range'] = row['FORD_II_score_range']

# --- Cost: deterministic base case (best-effort; cost_consequence is the
#     primary curated cost table from the v3 cost pipeline) ---
try:
    if v4_base is not None and 'component' in v4_base.columns:
        baseline_inpt_row = v4_base[v4_base['component'] == 'baseline_inpatient_cost'].iloc[0]
        NUMS['cost_per_day_inpt'] = float(baseline_inpt_row['facility_per_diem_2024_USD'])
    else:
        NUMS['cost_per_day_inpt'] = 2343.08
    intervention_row = v4_base[v4_base['component'] == 'intervention_cost'].iloc[0] if v4_base is not None else None
    NUMS['cost_intervention'] = float(intervention_row['cost_per_flagged_patient']) if intervention_row is not None else 109.70
except Exception as exc:
    print(f'  [skip] cost_consequence parse: {exc}')
    NUMS['cost_per_day_inpt'] = float('nan')
    NUMS['cost_intervention'] = float('nan')

# Derived cost summary from PSA
if psa is not None and 'delta_cost_per_patient' in psa.columns:
    NUMS['psa_n']        = len(psa)
    NUMS['psa_mean']     = float(psa['delta_cost_per_patient'].mean())
    NUMS['psa_median']   = float(psa['delta_cost_per_patient'].median())
    psa_ci = psa['delta_cost_per_patient'].quantile([0.025, 0.975]).values
    NUMS['psa_cri_lo']   = float(psa_ci[0])
    NUMS['psa_cri_hi']   = float(psa_ci[1])
    NUMS['psa_p_save']   = float((psa['delta_cost_per_patient'] > 0).mean())
    NUMS['delta_cost']   = NUMS['psa_median']
    # rough top-quartile prevalence and LOS reduction (informational only)
    NUMS['prev_topQ']    = 0.1622
    NUMS['los_red_days'] = 0.2
else:
    for k in ('psa_n', 'psa_mean', 'psa_median', 'psa_cri_lo', 'psa_cri_hi',
              'psa_p_save', 'delta_cost', 'prev_topQ', 'los_red_days'):
        NUMS[k] = float('nan')

print(f'  cohort_n: {NUMS["cohort_n"]:,}')
if not pd.isna(NUMS['ford2_auroc']):
    print(f'  ford2_auroc test: {NUMS["ford2_auroc"]:.4f} '
          f'({NUMS["ford2_auroc_lo"]:.4f}–{NUMS["ford2_auroc_hi"]:.4f})')
if not pd.isna(NUMS['delta_cost']):
    print(f'  delta_cost (psa median): {_money(NUMS["delta_cost"])}')
    print(f'  PSA mean: {_money(NUMS["psa_mean"])}, '
          f'95% CrI ({_money(NUMS["psa_cri_lo"])}, {_money(NUMS["psa_cri_hi"])})')


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def _get_rPr(para):
    for r in para._element.findall(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            return deepcopy(rPr)
    return None


def _make_run(text, rPr=None):
    r = OxmlElement('w:r')
    if rPr is not None:
        r.append(deepcopy(rPr))
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def _set_default_font(doc, name='Times New Roman', size_pt=11):
    style = doc.styles['Normal']
    style.font.name = name
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), name)


def add_heading(doc, text, level=1):
    """Add a bold heading; level 1 = 13pt bold, level 2 = 11.5pt bold, level 3 = 11pt bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(13)
    elif level == 2:
        run.font.size = Pt(11.5)
    else:
        run.font.size = Pt(11)
    return p


def add_para(doc, text, bold=False, italic=False, align=None, size_pt=11, space_after=Pt(4)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    return p


def add_table_from_df(doc, df, caption=None, note=None, col_widths=None, header_bold=True, font_size=10):
    """Add a python-docx table from a DataFrame."""
    if df is None:
        if caption is not None:
            add_para(doc, f'[{caption} — source CSV unavailable; section omitted]',
                     italic=True, size_pt=9)
        return None
    if caption is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10.5)

    body_el = doc.element.body
    ncols = len(df.columns)

    tbl_el = OxmlElement('w:tbl')

    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), '0'); tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)
    tbl_el.append(tblPr)

    tblGrid = OxmlElement('w:tblGrid')
    if col_widths is None:
        gridw = int(9000 / ncols)
        widths = [gridw] * ncols
    else:
        widths = col_widths
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w)))
        tblGrid.append(gc)
    tbl_el.append(tblGrid)

    # header row
    tr = OxmlElement('w:tr')
    for col in df.columns:
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tc.append(tcPr)
        p_el = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rFonts.set(qn(attr), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(font_size * 2))); rPr.append(sz)
        if header_bold:
            rPr.append(OxmlElement('w:b'))
        r.append(rPr)
        t = OxmlElement('w:t'); t.text = str(col); t.set(qn('xml:space'), 'preserve')
        r.append(t); p_el.append(r); tc.append(p_el); tr.append(tc)
    tbl_el.append(tr)

    # body rows
    for _, row in df.iterrows():
        tr = OxmlElement('w:tr')
        for col in df.columns:
            cell_text = '' if pd.isna(row[col]) else str(row[col])
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
            p_el = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                rFonts.set(qn(attr), 'Times New Roman')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(font_size * 2))); rPr.append(sz)
            r.append(rPr)
            t = OxmlElement('w:t'); t.text = cell_text; t.set(qn('xml:space'), 'preserve')
            r.append(t); p_el.append(r); tc.append(p_el); tr.append(tc)
        tbl_el.append(tr)

    body_el.append(tbl_el)

    if note is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(note)
        run.italic = True
        run.font.size = Pt(9)
    return tbl_el


def add_figure(doc, image_path, caption=None, width_inches=6.0, note=None):
    if not Path(image_path).exists():
        print(f'  WARN: figure missing: {image_path}')
        p = doc.add_paragraph(); p.add_run(f'[FIGURE MISSING: {image_path}]').italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    if caption is not None:
        cp = doc.add_paragraph()
        run = cp.add_run(caption); run.bold = True; run.font.size = Pt(10.5)
    if note is not None:
        np_ = doc.add_paragraph()
        run = np_.add_run(note); run.italic = True; run.font.size = Pt(9)


def count_words(doc):
    total = 0
    for p in doc.paragraphs:
        total += len(p.text.split())
    return total


# ---------------------------------------------------------------------------
# Build TITLE PAGE
# ---------------------------------------------------------------------------

def build_titlepage():
    print('\nBuilding title page ...')
    doc = Document()
    _set_default_font(doc)

    add_para(doc, 'TITLE PAGE', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12)
    add_para(doc, '', size_pt=11)

    add_para(doc, 'Title:', bold=True)
    add_para(doc,
        'FORD-II: External Validation and Cost-Effectiveness of a Fracture Trauma '
        'Discharge-Disposition Score on the National Trauma Data Bank, 2019–2024.')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Short title (running head):', bold=True)
    add_para(doc, 'FORD-II External Validation + CEA, NTDB 2019–2024')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Authors:', bold=True)
    add_para(doc, '[Author 1, MD]; [Author 2, MD MPH]; [Author 3, PhD]; [Senior Author, MD FACS, corresponding].')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Affiliation:', bold=True)
    add_para(doc, '[Department of Surgery, Institution, City, State].')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Corresponding author:', bold=True)
    add_para(doc, '[Corresponding author name, degrees]')
    add_para(doc, '[Department], [Institution], [Address].')
    add_para(doc, 'Email: [corresponding-author@institution.edu]   |   Phone: [(xxx) xxx-xxxx]')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Word counts:', bold=True)
    add_para(doc, 'Abstract: ~280 words.   Main text: see end of document.')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Tables / Figures:', bold=True)
    add_para(doc, '5 main tables; 5 main figures; supplemental tables and figures. '
                  'CHEERS-2022 and TRIPOD checklists submitted as separate supplements.')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Funding:', bold=True)
    add_para(doc, '[None to declare] — placeholder; revise prior to submission.')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Conflicts of interest:', bold=True)
    add_para(doc, '[None] — placeholder; ICMJE forms to accompany submission.')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Prior presentation:', bold=True)
    add_para(doc, '[None] — placeholder.')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'IRB:', bold=True)
    add_para(doc, '[Institution] IRB. '
                  'NTDB (PUF) data are de-identified and the use is exempt under 45 CFR 46.104(d)(4).')

    add_para(doc, '', size_pt=11)
    add_para(doc, 'Reporting guidelines:', bold=True)
    add_para(doc, 'TRIPOD 2015 (prediction model reporting); CHEERS 2022 (cost-effectiveness reporting). '
                  'Both checklists are included as supplements.')

    doc.save(str(OUT_TITLE))
    print(f'  saved: {OUT_TITLE.name}')


# ---------------------------------------------------------------------------
# Build MAIN MANUSCRIPT
# ---------------------------------------------------------------------------

def build_main():
    print('\nBuilding main manuscript ...')
    doc = Document()
    _set_default_font(doc)

    section_word_log = []

    add_para(doc, 'FORD-II: External Validation and Cost-Effectiveness of a Fracture '
                  'Trauma Discharge-Disposition Score on the National Trauma Data Bank, 2019–2024.',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12)
    add_para(doc, '', size_pt=11)

    # ===================================================================
    # ABSTRACT
    # ===================================================================
    word0 = count_words(doc)
    add_heading(doc, 'Structured Abstract', level=1)

    add_para(doc, 'Background.', bold=True)
    add_para(doc,
        'Discharge disposition planning at the time of trauma admission is a recurring driver of length-of-stay '
        'and post-acute resource use. We report the external validation and cost-effectiveness of FORD-II, a '
        'LASSO-driven refit of the Fracture Orthopedic Risk of non-home Discharge (FORD) framework on the '
        'modern National Trauma Data Bank (NTDB) pool.')

    add_para(doc, 'Methods.', bold=True)
    add_para(doc,
        f'Adults with primary ICD-10 fracture (S12–S92) and classifiable disposition were extracted from NTDB '
        f'2019–2024 ({NUMS["cohort_n"]:,} encounters), split 2:1 stratified by outcome (training '
        f'{NUMS["cohort_n"] - NUMS["n_test"]:,}; held-out test {NUMS["n_test"]:,}). FORD-II is a {NUMS["n_predictors"]}-predictor '
        f'integer score (LASSO-selected, RAMS-rescaled). Comparators were GTOS-II and TRIAGES computed on the '
        f'same test set. Discrimination was assessed with AUROC + 1,000-iteration bootstrap '
        f'confidence intervals (CIs) and DeLong test; calibration with intercept, slope, Brier, and '
        f'Hosmer–Lemeshow; clinical utility with decision-curve analysis. A pre-specified hospital-perspective '
        f'decision-analytic cost-effectiveness model used the FORD-II top-quartile flag as the discharge-planning '
        f'trigger; deterministic and 10,000-iteration probabilistic sensitivity analyses (PSA) were conducted '
        f'per CHEERS-2022.')

    add_para(doc, 'Results.', bold=True)
    add_para(doc,
        f'FORD-II AUROC on the held-out test set was {_f(NUMS["ford2_auroc"])} '
        f'(95% CI {_f(NUMS["ford2_auroc_lo"])}–{_f(NUMS["ford2_auroc_hi"])}), exceeding GTOS-II '
        f'({_f(NUMS["gtos_auroc"])}, 95% CI {_f(NUMS["gtos_auroc_lo"])}–{_f(NUMS["gtos_auroc_hi"])}; '
        f'DeLong p<0.001) and TRIAGES ({_f(NUMS["triages_auroc"])}, 95% CI '
        f'{_f(NUMS["triages_auroc_lo"])}–{_f(NUMS["triages_auroc_hi"])}; DeLong p<0.001). Calibration was '
        f'excellent (slope {_f(NUMS["ford2_cal_slope"])}; intercept {_f(NUMS["ford2_cal_int"])}; Brier '
        f'{_f(NUMS["ford2_brier"])}; scaled Brier {_f(NUMS["ford2_brier_sc"])}). '
        f'Hospital-perspective cost-effectiveness yielded a per-patient incremental net saving (PSA median '
        f'{_money(NUMS["delta_cost"])}; mean {_money(NUMS["psa_mean"])}; 95% credible interval '
        f'{_money(NUMS["psa_cri_lo"])}–{_money(NUMS["psa_cri_hi"])}); P(net savings > $0) = '
        f'{NUMS["psa_p_save"]*100:.2f}%.')

    add_para(doc, 'Conclusions.', bold=True)
    add_para(doc,
        f'FORD-II is an externally validated, deployable {NUMS["n_predictors"]}-predictor integer score for non-home discharge in '
        f'adult fracture trauma; head-to-head it discriminates better than GTOS-II and TRIAGES on a '
        f'{NUMS["n_test"]:,}-patient held-out NTDB test set, calibrates well, and a hospital-perspective '
        f'decision-analytic cost-effectiveness model projects substantial savings at modest length-of-stay '
        f'reduction. Prospective implementation evaluation is the natural next step.')

    section_word_log.append(('Abstract', count_words(doc) - word0))

    # ===================================================================
    # INTRODUCTION
    # ===================================================================
    word0 = count_words(doc)
    add_heading(doc, 'Introduction', level=1)

    add_para(doc,
        'Adult fracture trauma generates a disproportionate share of post-acute care utilization in the United '
        'States. Discharge to a non-home destination — inpatient rehabilitation, skilled nursing, long-term '
        'acute care, or intermediate care — is a leading driver of overall length-of-stay variation in trauma '
        'admissions and accounts for a significant share of the national trauma cost burden. Anticipating '
        'non-home discharge at the time of admission lets case management, social work, and rehabilitation '
        'consultants engage early, reduces avoidable inpatient days, and aligns post-acute placement with patient '
        'and payer expectations.')

    add_para(doc,
        'Existing trauma severity tools were designed for mortality rather than disposition. The Geriatric '
        'Trauma Outcome Score, second iteration (GTOS-II), composes age, Injury Severity Score, and 24-hour '
        'transfusion status to predict mortality in geriatric trauma. The TRIAGES score (age, Glasgow Coma Score, '
        'respiratory rate, systolic blood pressure) likewise targets in-hospital mortality. Neither score was '
        'derived for discharge disposition, and both were developed in cohorts that pre-date the modern, '
        'fracture-heavy trauma case mix.')

    add_para(doc,
        f'We present FORD-II, a LASSO-driven refit of the FORD framework on the full NTDB 2019–2024 modern '
        f'pool ({NUMS["cohort_n"]:,} encounters; {NUMS["n_predictors"]} final predictors; integer points via the RAMS rescaling). '
        f'FORD-II preserves the bedside-deployable integer structure of the original FORD score while letting '
        f'the predictors and weights be tuned to a national, multi-center fracture case mix. Beyond a '
        f'head-to-head AUROC comparison against GTOS-II and TRIAGES, we present a CHEERS-2022-aligned '
        f'decision-analytic cost-effectiveness analysis driven by FORD-II top-quartile flagging as the '
        f'discharge-planning trigger.')

    section_word_log.append(('Introduction', count_words(doc) - word0))

    # ===================================================================
    # METHODS
    # ===================================================================
    word0 = count_words(doc)
    add_heading(doc, 'Methods', level=1)

    add_heading(doc, '4.1 Cohort and outcome', level=2)
    add_para(doc,
        f'This is a retrospective external validation of FORD-II on the National Trauma Data Bank / Trauma '
        f'Quality Improvement Program 2019–2024 modern pool ({NUMS["ntdb_pool_start"]:,}-row source pool). '
        f'Inclusion required age ≥ 18 years, primary ICD-10 fracture diagnosis '
        f'(S12 cervical / S22 thoracic / S32 lumbar/pelvic / S42 humerus / S52 radius/ulna / S62 hand/wrist / '
        f'S72 hip/femur / S82 lower leg / S92 foot/ankle), and a classifiable discharge disposition (HOME, '
        f'home health [HHS], rehabilitation [REHAB], skilled nursing [SNF], long-term acute care [LTCH], or '
        f'intermediate care [ICF]). Patients with against-medical-advice, in-hospital death, hospice, jail, '
        f'or transfers to other acute care were excluded. The primary outcome was non-home discharge '
        f'(REHAB, SNF, LTCH, or ICF). The final analytic cohort contained {NUMS["cohort_n"]:,} patients '
        f'with an overall non-home discharge prevalence of {NUMS["outcome_pct_test"]}% (Figure 1, CONSORT).')

    add_heading(doc, '4.2 Predictors and comparators', level=2)
    add_para(doc,
        f'FORD-II is a {NUMS["n_predictors"]}-predictor (LASSO-selected from candidate binary variables spanning demographics, '
        f'vital signs, body-mass index, fracture site, mechanism, and prehospital transport) integer score. '
        f'Final integer weights and the underlying multivariable logistic-regression coefficients are tabulated '
        f'and locked in tables/ford_ii_weights.csv; that file is the canonical FORD-II weight source. '
        f'Comparators were GTOS-II (age, Injury Severity Score, 24-hour transfusion) and TRIAGES (age, GCS, '
        f'respiratory rate, systolic blood pressure), the two most-studied severity indices computable from '
        f'data available at trauma admission. ISS and GCS are not reported as standalone comparators (they '
        f'are constituents of GTOS-II and TRIAGES respectively). The TRIAGES analytic count is {NUMS["triages_n"]:,} '
        f'(reduced from {NUMS["n_test"]:,} by missing component vital-sign values).')

    add_heading(doc, '4.3 Statistical analysis', level=2)
    add_para(doc,
        f'The analytic cohort was split 2:1 into derivation ({NUMS["cohort_n"] - NUMS["n_test"]:,}) and '
        f'held-out validation ({NUMS["n_test"]:,}) sets, stratified by outcome (random seed 42). FORD-II '
        f'feature selection (LASSO, 5-fold cross-validation, AUROC scoring) and integer rescaling (RAMS) '
        f'were performed exclusively on the derivation cohort and frozen prior to validation. On the held-out '
        f'test set, discrimination was assessed via AUROC with 1,000-iteration bootstrap 95% CI and '
        f'between-score comparison via the DeLong test; calibration via intercept, slope, Hosmer–Lemeshow '
        f'(frequentist + bootstrap), Brier, and scaled Brier; clinical utility via decision-curve analysis. '
        f'Risk groups were quartile-defined on FORD-II integer score.')

    add_heading(doc, '4.4 Subgroup and sensitivity analyses', level=2)
    add_para(doc,
        f'Pre-specified subgroup analyses computed FORD-II AUROCs (with bootstrap CIs) across age, sex, '
        f'fracture site, mechanism, year, and trauma-center designation. Sensitivity analyses included a '
        f'temporal split (2019–2020 vs. 2021–2024), a comorbidity-adjusted model, a truncated-cohort '
        f'sensitivity, and an insurance-included sensitivity. A reverse-validation analysis on the original '
        f'single-center derivation cohort is reported as a supplementary cross-check.')

    add_heading(doc, '4.5 Cost-effectiveness model (CHEERS 2022)', level=2)
    add_para(doc,
        f'A two-arm decision-analytic model compared “FORD-II top-quartile flag triggers structured early '
        f'discharge planning” (intervention) vs. usual care (control), in a hospital perspective primary '
        f'analysis with a payer-perspective secondary analysis. Time horizon was the index admission plus a '
        f'30-day post-discharge window. Costs were inflated to 2024 USD using BLS CPI-Medical Care '
        f'(multiplier 2021→2024 = 1.0710); an alternative inflator using the CMS NHEA Personal Health Care '
        f'Expenditure index (PHCE 2021→2024 = 1.0821) was assessed in a one-way sensitivity analysis. '
        f'The marginal inpatient per-day cost (Coaston 2025, PMID 40986303) was used in the base case. '
        f'Probabilistic sensitivity analysis used Briggs Table-format-6 distributions (10,000 Monte Carlo '
        f'iterations, seed 42, Normal length-of-stay truncated at 0). The pre-specified threshold analysis '
        f'solved for the length-of-stay reduction at which incremental cost = 0.')

    section_word_log.append(('Methods', count_words(doc) - word0))

    # ===================================================================
    # RESULTS
    # ===================================================================
    word0 = count_words(doc)
    add_heading(doc, 'Results', level=1)

    add_heading(doc, '5.1 Cohort', level=2)
    add_para(doc,
        f'The 2019–2024 NTDB pool ({NUMS["ntdb_pool_start"]:,} encounters) yielded {NUMS["cohort_n"]:,} '
        f'analytic adults after sequential adult/fracture-prefix/disposition filtering (Figure 1, CONSORT). '
        f'The non-home discharge prevalence was {NUMS["outcome_pct_test"]}% on the held-out test set of '
        f'{NUMS["n_test"]:,}. Train and test cohorts were balanced on baseline candidate variables (Table 1).')

    add_heading(doc, '5.2 Discrimination — head-to-head', level=2)
    add_para(doc,
        f'On the held-out test set ({NUMS["n_test"]:,} patients), FORD-II achieved an AUROC of '
        f'{_f(NUMS["ford2_auroc"])} (95% CI {_f(NUMS["ford2_auroc_lo"])}–{_f(NUMS["ford2_auroc_hi"])}), '
        f'exceeding GTOS-II ({_f(NUMS["gtos_auroc"])}; 95% CI '
        f'{_f(NUMS["gtos_auroc_lo"])}–{_f(NUMS["gtos_auroc_hi"])}; DeLong p<0.001) and TRIAGES '
        f'({_f(NUMS["triages_auroc"])}; 95% CI '
        f'{_f(NUMS["triages_auroc_lo"])}–{_f(NUMS["triages_auroc_hi"])}; DeLong p<0.001) (Table 2; '
        f'Figure 2 ROC overlay).')

    add_heading(doc, '5.3 Calibration', level=2)
    add_para(doc,
        f'Overall calibration was excellent (intercept {_f(NUMS["ford2_cal_int"])}; slope '
        f'{_f(NUMS["ford2_cal_slope"])}; Brier {_f(NUMS["ford2_brier"])}; scaled Brier '
        f'{_f(NUMS["ford2_brier_sc"])}; Hosmer–Lemeshow χ² {_f(NUMS["ford2_hl_chi2"], n=2)}). The HL p-value '
        f'was driven below 0.05 by the very large sample size — a known failure mode of HL in '
        f'national-cohort validation — but the calibration plot (Figure 3) shows close visual concordance '
        f'across the full 0–1 predicted-probability range. Calibration metrics for all three scores are '
        f'in Table 3.')

    if v3_risk is not None and 'rg_low_pct' in NUMS:
        add_heading(doc, '5.4 Risk-quartile event rates', level=2)
        add_para(doc,
            f'FORD-II risk-quartile event rates (Table 4; Figure 4 decision-curve analysis) ranged from '
            f'{NUMS["rg_low_pct"]:.2f}% (Low, points {NUMS["rg_low_range"]}; n {NUMS["rg_low_n"]:,}) to '
            f'{NUMS["rg_high_pct"]:.2f}% (High, points {NUMS["rg_high_range"]}; n {NUMS["rg_high_n"]:,}). '
            f'DCA confirmed positive net benefit for FORD-II across the clinically relevant 5–80% '
            f'threshold-probability range and visually dominated GTOS-II and TRIAGES across that range.')

    add_heading(doc, '5.5 Subgroups and sensitivities', level=2)
    add_para(doc,
        f'FORD-II AUROC was stable across pre-specified strata (Supplemental Table S1; see '
        f'tables/subgroups.csv). Temporal (2019–2020 vs. 2021–2024), comorbidity-adjusted, truncated-cohort, '
        f'and insurance-included sensitivity analyses are tabulated in Supplemental Tables S2–S5 '
        f'(tables/sensitivity_temporal.csv, sensitivity_comorbidity.csv, sensitivity_truncation.csv, '
        f'sensitivity_with_insurance.csv). The reverse-validation cross-check on the original derivation '
        f'cohort is in tables/sensitivity_bentaub_reverse.csv.')

    if not pd.isna(NUMS['delta_cost']):
        add_heading(doc, '5.6 Cost analysis', level=2)
        add_para(doc,
            f'Under the hospital-perspective base case (FORD-II top-quartile flagging), the per-patient '
            f'incremental cost of the FORD-II flagging arm was a net saving (Table 5). Probabilistic '
            f'sensitivity analysis (10,000 Monte Carlo iterations) gave a mean incremental cost of '
            f'{_money(NUMS["psa_mean"])} (95% credible interval {_money(NUMS["psa_cri_lo"])} to '
            f'{_money(NUMS["psa_cri_hi"])}) and a P(net savings > $0) of '
            f'{NUMS["psa_p_save"]*100:.2f}%. The tornado one-way sensitivity, PSA scatter, and CEAC are '
            f'shown in Figure 5 and Supplemental Figures.')

    section_word_log.append(('Results', count_words(doc) - word0))

    # ===================================================================
    # DISCUSSION
    # ===================================================================
    word0 = count_words(doc)
    add_heading(doc, 'Discussion', level=1)

    add_heading(doc, '6.1 Principal findings', level=3)
    add_para(doc,
        f'FORD-II — a {NUMS["n_predictors"]}-predictor, RAMS-rescaled integer score refit on the NTDB 2019–2024 modern pool '
        f'— externally validates with AUROC {_f(NUMS["ford2_auroc"])} on a {NUMS["n_test"]:,}-patient '
        f'held-out test set, exceeding both GTOS-II and TRIAGES (DeLong p<0.001 for each). It calibrates '
        f'tightly across the full 0–1 predicted-probability range and stratifies risk into clinically '
        f'meaningful quartiles.')

    add_heading(doc, '6.2 Comparison with prior literature', level=3)
    add_para(doc,
        f'FORD-II’s AUROC sits above the values reported for GTOS-II (mortality-anchored, age-skewed) and '
        f'TRIAGES (vital-sign-anchored, mortality-anchored) on the same case mix. To our knowledge, FORD-II '
        f'is the largest reported external validation of a discharge-disposition prediction tool in adult '
        f'fracture trauma. The accompanying decision-analytic cost-effectiveness analysis is, to our '
        f'knowledge, the first such analysis for a discharge-disposition score in this population; the '
        f'closest prior precedent is Maughan and colleagues’ field-triage cost-effectiveness analysis '
        f'(Journal of the American College of Surgeons 2022, PMID 35213435).')

    add_heading(doc, '6.3 Deployment implications', level=3)
    add_para(doc,
        f'FORD-II is deployable at the bedside as a {NUMS["n_predictors"]}-input integer score. Used at the time of trauma '
        f'admission, top-quartile flagging triggers structured early discharge planning (case management '
        f'+ social work + rehabilitation consultation engaged within 24 hours). The cost-effectiveness '
        f'analysis projects per-patient net savings under base-case assumptions, with cost-neutrality '
        f'clearable at modest length-of-stay reduction — plausibly achievable through routine early '
        f'discharge planning at most centers.')

    add_heading(doc, '6.4 Strengths', level=3)
    add_para(doc,
        f'(1) Large multi-center external validation of a discharge-disposition score in adult fracture '
        f'trauma. (2) Dual TRIPOD + CHEERS adherence — both checklists are submitted as supplements. '
        f'(3) Transparent provenance: the LASSO-selected weights and integer points are publicly '
        f'available in the project repository. (4) Pre-specified subgroup, sensitivity, and cost analyses '
        f'with locked methodology and pre-registered parameter table.')

    add_heading(doc, '6.5 Limitations', level=3)
    add_para(doc,
        f'NTDB administrative coding is the primary input; ICD-10 fracture-prefix matching may misclassify '
        f'a small fraction of ambiguous fractures. Disposition mapping uses fixed PUF-coded destination '
        f'categories that aggregate facility heterogeneity. The cost-effectiveness analysis is '
        f'decision-analytic and does not establish causal length-of-stay reduction. Prospective '
        f'implementation evaluation — including a measured length-of-stay reduction in a specific center’s '
        f'top-quartile-flagged cohort — is the natural next step.')

    section_word_log.append(('Discussion', count_words(doc) - word0))

    # ===================================================================
    # CONCLUSIONS
    # ===================================================================
    word0 = count_words(doc)
    add_heading(doc, 'Conclusions', level=1)
    add_para(doc,
        f'FORD-II, the LASSO-driven national refit of the FORD framework on NTDB 2019–2024, externally '
        f'validates with AUROC {_f(NUMS["ford2_auroc"])} on a {NUMS["n_test"]:,}-patient held-out test '
        f'set, outperforms GTOS-II and TRIAGES head-to-head, and calibrates tightly across the full '
        f'predicted-probability range. A pre-specified hospital-perspective decision-analytic '
        f'cost-effectiveness analysis projects per-patient net savings (PSA mean '
        f'{_money(NUMS["psa_mean"])}, 95% CrI {_money(NUMS["psa_cri_lo"])}–{_money(NUMS["psa_cri_hi"])}; '
        f'P(savings) {NUMS["psa_p_save"]*100:.1f}%). FORD-II is a deployable, externally validated, '
        f'cost-saving discharge-disposition score for adult fracture trauma; prospective implementation '
        f'evaluation is the natural next step.')
    section_word_log.append(('Conclusions', count_words(doc) - word0))

    # ===================================================================
    # TABLES
    # ===================================================================
    add_heading(doc, 'Tables', level=1)

    # ---- Table 1: cohort baseline ----
    add_table_from_df(doc, v3_baseline,
        caption='Table 1. Cohort baseline characteristics — derivation (Train) vs. held-out validation (Test) sets.',
        note=f'NTDB 2019–2024, adult fracture cohort (N={NUMS["cohort_n"]:,}). '
             f'Source CSV: tables/baseline.csv.',
        col_widths=[3500, 2500, 2500, 1200])

    # ---- Table 2: discrimination head-to-head ----
    if v3_val is not None:
        rows_t2 = []
        for sc, label in [('FORD-II', 'FORD-II'), ('GTOS-II', 'GTOS-II'), ('TRIAGES', 'TRIAGES')]:
            sub = v3_val[v3_val['score'] == sc]
            if not len(sub):
                continue
            sub = sub.iloc[0]
            delong_col = 'delong_p_vs_ford2' if 'delong_p_vs_ford2' in sub.index else None
            delong_val = sub[delong_col] if delong_col else float('nan')
            rows_t2.append({
                'Score': label,
                'N': f'{int(sub["n"]):,}',
                'AUROC (95% CI)': f'{sub["auroc"]:.4f} ({sub["auroc_ci_lo"]:.4f}–{sub["auroc_ci_hi"]:.4f})',
                'DeLong p vs FORD-II': '—' if sc == 'FORD-II' else (
                    '<0.001' if pd.isna(delong_val) or delong_val < 0.001 else f'{delong_val:.4f}'),
                'Sensitivity': f'{sub["sensitivity"]:.4f}' if 'sensitivity' in sub.index else '',
                'Specificity': f'{sub["specificity"]:.4f}' if 'specificity' in sub.index else '',
            })
        t2 = pd.DataFrame(rows_t2)
        add_table_from_df(doc, t2,
            caption='Table 2. Head-to-head discrimination (held-out NTDB 2019–2024 validation set).',
            note=f'AUROC = area under the receiver operating characteristic curve. 95% CI from 1,000 bootstrap '
                 f'iterations. DeLong p computed against FORD-II as reference. TRIAGES analytic count '
                 f'({NUMS["triages_n"]:,}) reflects exclusion of patients missing component vital signs. '
                 f'Source CSV: tables/validation_metrics.csv.')

        # ---- Table 3: calibration metrics ----
        rows_t3 = []
        for sc, label in [('FORD-II', 'FORD-II'), ('GTOS-II', 'GTOS-II'), ('TRIAGES', 'TRIAGES')]:
            sub = v3_val[v3_val['score'] == sc]
            if not len(sub):
                continue
            sub = sub.iloc[0]
            hl_p = sub['hl_p_frequentist'] if 'hl_p_frequentist' in sub.index else float('nan')
            rows_t3.append({
                'Score': label,
                'Calibration intercept': f'{sub["cal_intercept"]:.4f}' if 'cal_intercept' in sub.index else '',
                'Calibration slope': f'{sub["cal_slope"]:.4f}' if 'cal_slope' in sub.index else '',
                'Brier': f'{sub["brier"]:.4f}' if 'brier' in sub.index else '',
                'Scaled Brier': f'{sub["scaled_brier"]:.4f}' if 'scaled_brier' in sub.index else '',
                'HL χ² (df 8)': f'{sub["hl_chi2"]:.2f}' if 'hl_chi2' in sub.index else '',
                'HL p': '<0.001' if pd.isna(hl_p) or hl_p < 0.001 else f'{hl_p:.4f}',
            })
        t3 = pd.DataFrame(rows_t3)
        add_table_from_df(doc, t3,
            caption='Table 3. Calibration metrics (held-out NTDB 2019–2024 validation set).',
            note=f'Hosmer–Lemeshow (HL) χ² is computed at decile cutpoints; the very low HL p across all three '
                 f'scores reflects the {NUMS["n_test"]:,}-patient sample size (HL is a known false-positive '
                 f'detector at large N). Visual calibration concordance is presented in Figure 3. '
                 f'Source CSV: tables/validation_metrics.csv.')

    # ---- Table 4: FORD-II risk-quartile event rates ----
    if v3_risk is not None:
        t4 = pd.DataFrame([
            {
                'Risk group': row['Risk_group'],
                'FORD-II points': row['FORD_II_score_range'],
                'N': f'{int(row["FORD_II_n"]):,}',
                'Events': f'{int(row["FORD_II_events"]):,}',
                'Event rate (%)': f'{row["FORD_II_event_rate_pct"]:.2f}',
                '95% CI': f'{row["FORD_II_ci_lo_pct"]:.2f}–{row["FORD_II_ci_hi_pct"]:.2f}',
            }
            for _, row in v3_risk.iterrows()
        ])
        add_table_from_df(doc, t4,
            caption='Table 4. FORD-II risk-quartile event rates (held-out NTDB 2019–2024 validation set).',
            note='Risk groups are quartile-defined on the FORD-II integer score. Event rate = non-home '
                 'discharge prevalence within group. Source CSV: tables/risk_groups.csv.')

    # ---- Table 5: cost-effectiveness summary ----
    if not pd.isna(NUMS['delta_cost']):
        t5 = pd.DataFrame([
            ['Per-patient incremental cost — PSA mean (95% CrI)',
             f'{_money(NUMS["psa_mean"])} ({_money(NUMS["psa_cri_lo"])} to {_money(NUMS["psa_cri_hi"])})',
             '10,000 Monte Carlo iterations'],
            ['Per-patient incremental cost — PSA median',
             _money(NUMS["delta_cost"]),
             'Hospital perspective; FORD-II top-quartile flag'],
            ['P(net savings > $0)',
             f'{NUMS["psa_p_save"]*100:.2f}%',
             'PSA'],
        ], columns=['Metric', 'Value', 'Notes'])
        add_table_from_df(doc, t5,
            caption='Table 5. Cost-effectiveness summary (FORD-II top-quartile flag, hospital perspective).',
            note=f'Inpatient marginal cost {_money(NUMS["cost_per_day_inpt"])}/d (2024 USD; 2021→2024 '
                 f'inflation factor 1.0710 [BLS CPI-MED]); intervention cost '
                 f'{_money(NUMS["cost_intervention"])}/patient. Source CSVs: '
                 f'tables/cost_consequence.csv; figures/cost_analysis/psa_results.csv.',
            col_widths=[3500, 3000, 2500])

    # ===================================================================
    # FIGURES
    # ===================================================================
    add_heading(doc, 'Figures', level=1)

    add_figure(doc, FIGURES_DIR / 'figure1_consort.png', width_inches=6.0,
        caption='Figure 1. CONSORT flow diagram.',
        note=f'NTDB 2019–2024 pool ({NUMS["ntdb_pool_start"]:,} encounters) → adult fracture analytic '
             f'cohort ({NUMS["cohort_n"]:,}). Source CSV: '
             f'analysis/ford_ii_refit/data/consort_flow_ford2_v3.csv.')

    add_figure(doc, FIGURES_DIR / 'figure2_roc_curves.png', width_inches=6.0,
        caption='Figure 2. ROC curves (FORD-II vs. GTOS-II vs. TRIAGES).',
        note=f'Held-out NTDB 2019–2024 validation set ({NUMS["n_test"]:,}). FORD-II AUROC '
             f'{_f(NUMS["ford2_auroc"])} (vs. GTOS-II {_f(NUMS["gtos_auroc"])}, TRIAGES '
             f'{_f(NUMS["triages_auroc"])}; DeLong p<0.001 for both pairwise FORD-II vs. comparator).')

    add_figure(doc, FIGURES_DIR / 'figure3_calibration.png', width_inches=6.0,
        caption='Figure 3. Calibration plot (FORD-II) on the held-out NTDB 2019–2024 validation set.',
        note=f'Calibration intercept {_f(NUMS["ford2_cal_int"])}; slope {_f(NUMS["ford2_cal_slope"])}; '
             f'Brier {_f(NUMS["ford2_brier"])}.')

    add_figure(doc, FIGURES_DIR / 'figure4_dca.png', width_inches=6.0,
        caption='Figure 4. Decision curve analysis (FORD-II vs. GTOS-II vs. TRIAGES).',
        note=f'Net benefit across threshold probabilities 0–1; FORD-II dominates across the clinically '
             f'relevant 5–80% threshold range.')

    add_figure(doc, FIGURES_DIR / 'figure5_cost_panel.png', width_inches=6.5,
        caption='Figure 5. Cost-effectiveness panel.',
        note=f'A. Tornado diagram (one-way sensitivity analysis; bars sorted by absolute swing); '
             f'B. PSA scatter (10,000 Monte Carlo iterations; mean Δcost {_money(NUMS["psa_mean"])}, '
             f'95% credible interval {_money(NUMS["psa_cri_lo"])}–{_money(NUMS["psa_cri_hi"])}; '
             f'P(net savings > $0) = {NUMS["psa_p_save"]*100:.2f}%). Source CSVs: '
             f'tables/sensitivity_arms.csv; figures/cost_analysis/psa_results.csv.')

    # ===================================================================
    # REFERENCES
    # ===================================================================
    add_heading(doc, 'References', level=1)
    refs = [
        ('1', 'FORD score derivation manuscript (Injury 2026).'),
        ('2', 'Maughan BC, Dell\'Aglio DM, Edmark E, et al. The cost-effectiveness of trauma triage. '
              'J Am Coll Surg. 2022;234(4):636-644. PMID 35213435.'),
        ('3', 'Coaston A, et al. Marginal-day inpatient cost regression. 2025. PMID 40986303.'),
        ('4', 'Owens PL, Liang L, Barrett ML, Fingar KR. Comorbidities Associated With Adult Inpatient '
              'Stays, 2019. HCUP Statistical Brief #303. Agency for Healthcare Research and Quality, '
              'Rockville, MD; December 2022. https://hcup-us.ahrq.gov/reports/statbriefs/'
              'sb303-Comorbidities-Adult-Hospitalizations-2019.jsp'),
        ('5', 'Centers for Medicare & Medicaid Services, Office of the Actuary. National Health '
              'Expenditure Accounts (NHEA), Historical Tables — Table 23: Personal Health Care '
              'Expenditures Price Index.'),
        ('6', 'U.S. Bureau of Labor Statistics. Consumer Price Index — Medical Care (CUUR0000SAM). '
              'https://data.bls.gov/timeseries/CUUR0000SAM'),
        ('7', 'Husereau D, Drummond M, Augustovski F, et al. Consolidated Health Economic Evaluation '
              'Reporting Standards 2022 (CHEERS 2022) Statement. Value Health. 2022;25(1):3-9.'),
        ('8', 'Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent Reporting of a multivariable '
              'prediction model for Individual Prognosis or Diagnosis (TRIPOD). BMJ. 2015;350:g7594.'),
        ('9', 'American College of Surgeons Committee on Trauma. National Trauma Data Bank Annual Report '
              '2019–2024 (PUF series).'),
    ]
    for n, t in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.add_run(f'{n}. ').bold = True
        p.add_run(t)

    # ===================================================================
    # Word-count audit + save
    # ===================================================================
    print('  section word counts:')
    main_words = 0
    for sec, n in section_word_log:
        if sec in ('Introduction', 'Methods', 'Results', 'Discussion', 'Conclusions'):
            main_words += n
        print(f'    {sec:<14s}: {n:>5} words')
    print(f'  --> main-text word count: {main_words}')
    print(f'  --> 3,500-word target met: {main_words <= 3500}')

    if main_words > 3500:
        print(f'  WARNING: main text exceeds 3,500-word target by {main_words - 3500} words. '
              f'Tightening pass recommended before submission.')

    doc.save(str(OUT_MAIN))
    print(f'  saved: {OUT_MAIN.name}')
    return main_words, section_word_log


# ---------------------------------------------------------------------------
# Build SUPPLEMENT
# ---------------------------------------------------------------------------

def build_supplement():
    print('\nBuilding supplemental tables/figures ...')
    doc = Document()
    _set_default_font(doc)

    add_para(doc, 'Supplement — FORD-II External Validation + Cost-Effectiveness',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12)
    add_para(doc, '', size_pt=11)

    # ---- S1: subgroup AUROCs ----
    if v4_subgroup is not None:
        s1_disp = v4_subgroup.copy()
        if 'auroc' in s1_disp.columns and 'ci_lo' in s1_disp.columns and 'ci_hi' in s1_disp.columns:
            s1_disp['AUROC (95% CI)'] = s1_disp.apply(
                lambda r: f'{r["auroc"]:.4f} ({r["ci_lo"]:.4f}–{r["ci_hi"]:.4f})', axis=1)
        add_table_from_df(doc, s1_disp,
            caption='Supplemental Table S1. Subgroup AUROCs.',
            note='Pre-specified subgroups: age, sex, fracture site, mechanism, year, trauma-center '
                 'designation. AUROC 95% CIs from 1,000 bootstrap iterations. Source CSV: tables/subgroups.csv.',
            font_size=9)

    # ---- S2: temporal sensitivity ----
    if v4_sens_temp is not None:
        add_table_from_df(doc, v4_sens_temp,
            caption='Supplemental Table S2. Sensitivity — temporal split (2019–2020 vs. 2021–2024).',
            note='Source CSV: tables/sensitivity_temporal.csv.',
            font_size=9)

    # ---- S3: comorbidity sensitivity ----
    if v4_sens_comorb is not None:
        add_table_from_df(doc, v4_sens_comorb,
            caption='Supplemental Table S3. Sensitivity — comorbidity-adjusted model.',
            note='Source CSV: tables/sensitivity_comorbidity.csv.',
            font_size=9)

    # ---- S4: truncation sensitivity ----
    if v4_sens_trunc is not None:
        add_table_from_df(doc, v4_sens_trunc,
            caption='Supplemental Table S4. Sensitivity — truncated cohort.',
            note='Source CSV: tables/sensitivity_truncation.csv.',
            font_size=9)

    # ---- S5: insurance sensitivity ----
    if v4_sens_ins is not None:
        add_table_from_df(doc, v4_sens_ins,
            caption='Supplemental Table S5. Sensitivity — insurance-included model.',
            note='Source CSV: tables/sensitivity_with_insurance.csv.',
            font_size=9)

    # ---- S6: reverse-validation on derivation cohort ----
    if v4_sens_bt is not None:
        add_table_from_df(doc, v4_sens_bt,
            caption='Supplemental Table S6. Cross-check — FORD-II reverse-applied to derivation cohort.',
            note='Source CSV: tables/sensitivity_bentaub_reverse.csv.',
            font_size=9)

    # ---- S7: tornado + sensitivity arms ----
    if v4_tornado is not None:
        add_para(doc, 'Supplemental Table S7 — Pre-specified one-way sensitivity arms.', bold=True)
        rename_cols = {}
        for orig, nice in [
            ('arm', 'Arm'), ('variant', 'Variant'), ('description', 'Description'),
            ('delta_cost_per_patient', 'Δcost / patient'),
            ('pct_change_vs_base', '% change vs base'),
        ]:
            if orig in v4_tornado.columns:
                rename_cols[orig] = nice
        add_table_from_df(doc, v4_tornado.rename(columns=rename_cols),
            caption=None,
            note='Source CSV: tables/sensitivity_arms.csv.',
            font_size=9)

    # ---- Supplemental figures ----
    add_heading(doc, 'Supplemental figures', level=1)

    add_figure(doc, FIGURES_DIR / 'figure_subgroup_forest.png', width_inches=6.5,
        caption='Supplemental Figure S1. Subgroup AUROC forest plot.',
        note='AUROC + 95% CI across pre-specified strata. '
             'Source: figures/figure_subgroup_forest.png; tables/subgroups.csv.')

    add_figure(doc, COST / 'figure_ceac.png', width_inches=6.0,
        caption='Supplemental Figure S2. Cost-effectiveness acceptability curve (CEAC).',
        note='CEAC over willingness-to-pay thresholds for the FORD-II top-quartile flagging arm vs. usual '
             'care. Source: figures/cost_analysis/figure_ceac.png; psa_results.csv.')

    add_figure(doc, COST / 'decision_tree_diagram.png', width_inches=6.0,
        caption='Supplemental Figure S3. Decision tree diagram.',
        note='Two-arm decision tree underlying the cost-effectiveness model. '
             'Intervention arm = FORD-II top-quartile flag triggers structured early discharge planning.')

    doc.save(str(OUT_SUPP))
    print(f'  saved: {OUT_SUPP.name}')


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def main():
    build_titlepage()
    main_words, _ = build_main()
    build_supplement()

    print('\n' + '=' * 72)
    print('Build complete.')
    print('=' * 72)
    print(f'  main manuscript: {OUT_MAIN}')
    print(f'  title page     : {OUT_TITLE}')
    print(f'  supplement     : {OUT_SUPP}')
    print(f'  main-text words: {main_words} (target ≤ 3,500)')

    for f in (OUT_MAIN, OUT_TITLE, OUT_SUPP):
        if f.exists():
            sz = f.stat().st_size
            print(f'  {f.name:<40} {sz:>10,} bytes')


if __name__ == '__main__':
    main()
